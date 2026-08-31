"""Export the Akasha Platform data schema to Excel and Word.

Generates Akasha_Data_Schema.xlsx and Akasha_Data_Schema.docx in the project root.
The schema content below is maintained by hand and mirrors backend/models.py plus
the ingestion scripts. Regenerate whenever a model or ingestion script changes.

    python scripts/export_schema_docs.py
"""

import os
import sys
from datetime import datetime

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


GENERATED = "12 August 2026"
ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Palette shared with the HTML reference
COPPER = "9E5228"
TEAL = "2C6A73"
INK = "131B22"
SURFACE = "F1F4F7"
LINE = "DAE1E8"


SUMMARY = [
    ("Source systems", "6"),
    ("Tables", "27"),
    ("Rows stored", "613,336"),
    ("Mapped projects", "64"),
    ("Join keys", "3"),
]

JOIN_KEYS = [
    ("project_id", "p6_project.project_id",
     "All Primavera P6 tables - schedule, activities, WBS, baselines, resource assignments."),
    ("spv_plant_code", "plant_code",
     "SAP tables: mt_inventory, mt_poamount, mt_materialdocument, mt_slr_data."),
    ("module_wbs", "wbs_element (prefix match)",
     "Same four SAP tables. The first 6 characters of a WBS element identify the plant; "
     "deeper segments identify block and package."),
    ("id", "tc_*.mapping_id",
     "Transmission tables tc_project_entry and tc_network_edge (foreign key)."),
    ("project_name_from_p6", "Pulse / E-Invoice project names",
     "Quality and invoice records arrive with portal-side project names, matched back through this label."),
]

RETRIEVAL_RULE = (
    "Retrieval rule: resolve the project through project_mapping first, then query SAP by "
    "plant_code + wbs_element, or transmission by mapping_id. Fuzzy-matching project names "
    "directly against SAP or TC tables produces confident wrong answers."
)


def T(name, tag, rows, note, cols=None, callout=None):
    return {"name": name, "tag": tag, "rows": rows, "note": note,
            "cols": cols or [], "callout": callout}


DOMAINS = [
    {
        "key": "Project master",
        "blurb": "The single registry of projects the platform reports on, and the anchor for "
                 "every cross-system lookup.",
        "source": "Data/NEW31/AKASHA SAP MASTER FILE.xlsx -> scripts/ingest_mapping.py",
        "tables": [
            T("project_mapping", "Excel master", 64,
              "One row per SPV / plot. Holds capacity, portfolio grouping and all three join keys.",
              [
                  ("id", "int PK", "Internal key. Referenced by transmission tables. [TC KEY]"),
                  ("project", "text", "Project name as written in the master file."),
                  ("spv_name", "text", "Special-purpose vehicle (legal entity) owning the asset."),
                  ("project_id", "text", "Primavera P6 project identifier. [P6 KEY]"),
                  ("project_name_from_p6", "text", "Project name exactly as it appears in P6 - used for label matching against portal data."),
                  ("plot_no", "text", "Land plot number at the site."),
                  ("category", "text", "Asset category (solar, BESS, and so on)."),
                  ("mms_type", "text", "Module mounting structure type."),
                  ("capacity_mwac", "numeric", "Nameplate AC capacity in MW. Denominator for MW-based progress."),
                  ("capacity_mwdc", "numeric", "Nameplate DC capacity in MW."),
                  ("ol", "text", "Operating licence / ownership label from the master file."),
                  ("spv_plant_code", "text", "SAP plant code. [SAP KEY 1]"),
                  ("module_wbs", "text", "SAP WBS element prefix for the module scope. [SAP KEY 2]"),
                  ("agel / age6l", "text", "Group entity codes carried from the master file."),
                  ("cluster", "text", "Geographic portfolio grouping (for example Khavda, Rajasthan)."),
                  ("subcluster", "text", "Sub-cluster or EPC contractor."),
                  ("not_allocated", "text", "Flag for scope not yet allocated to an SPV."),
                  ("source_of_origin", "text", "Which master sheet or revision the row came from."),
                  ("priority", "text", "Business priority tier."),
                  ("is_commissioned", "bool", "Whether the project has reached commissioning."),
                  ("tc_progress", "jsonb", "Cached transmission progress snapshot. Cheap to read but can go stale - check tc_network_edge for anything time-sensitive."),
              ]),
        ],
    },
    {
        "key": "Schedule & progress (P6)",
        "blurb": "The construction schedule: what is planned, what is done, what has slipped, and "
                 "where the critical path runs. The largest domain by volume.",
        "source": "Oracle Primavera P6 Cloud REST API - /project, /baselineProject, /wbs, /activity, "
                  "/resourceAssignment, /activityRisk - services/p6_service.py",
        "tables": [
            T("p6_project", "Live API", 67,
              "Project-level rollups - one row per P6 project. Every headline schedule and cost KPI "
              "on the dashboard originates here.",
              [
                  ("p6_object_id", "bigint UQ", "P6 internal object ID. Foreign key target for activities, baselines and assignments."),
                  ("project_id", "text", "P6 'Id' field - the canonical project identifier. [P6 KEY]"),
                  ("name", "text", "Project name in P6."),
                  ("status", "text", "Active / Planned / Inactive. Determines whether the other metrics are meaningful at all."),
                  ("start_date, finish_date", "timestamp", "Current schedule start and finish."),
                  ("planned_start_date, scheduled_finish_date", "timestamp", "Planned start and the finish the current schedule calculates."),
                  ("data_date", "timestamp", "Schedule cutoff - progress after this date is not reflected. Drives the freshness indicator."),
                  ("must_finish_by_date", "timestamp", "Contractual deadline, independent of the schedule. Negative slack against this is a breach risk even when SPI looks healthy."),
                  ("duration_percent_complete", "numeric", "Duration-weighted percent complete."),
                  ("planned_duration, actual_duration, remaining_duration", "numeric", "Duration rollups in days."),
                  ("actual_non_labor_units, baseline_non_labor_units", "numeric", "Units-based progress (actual vs baseline)."),
                  ("budget_labor_units, at_completion_non_labor_units", "numeric", "Budget at completion and forecast units at completion."),
                  ("activity_count", "int", "Total activities in the project."),
                  ("completed_ / in_progress_ / not_started_activity_count", "int", "Activity status rollup. A large not-started bucket late in the schedule is a risk signal."),
                  ("total_float", "numeric", "Schedule buffer in days on the driving path. Zero or below means critical - no buffer left."),
                  ("finish_date_variance, start_date_variance", "numeric", "Days of deviation from baseline. Positive means late."),
                  ("duration_variance", "numeric", "Duration deviation from baseline."),
                  ("actual_total_cost, planned_cost", "numeric", "Spend to date and budget."),
                  ("cost_performance_index", "numeric", "CPI - earned value cost efficiency. Below 1.0 means over budget for work done."),
                  ("schedule_performance_index", "numeric", "SPI - earned value schedule efficiency. Below 1.0 means behind plan."),
                  ("current_budget, total_cost_variance", "numeric", "Approved budget and variance against it."),
                  ("location_name, parent_eps_name", "text", "Site and position in the P6 enterprise project structure."),
                  ("baseline_start_date, baseline_finish_date, baseline_duration, baseline_total_cost", "timestamp / numeric", "The originally approved plan, summarised on the project row."),
                  ("baseline_completed_ / in_progress_ / not_started_activity_count", "int", "Baseline activity status rollup."),
                  ("current_baseline_project_object_id", "bigint", "Which baseline is currently active."),
                  ("last_synced_at", "timestamp", "When the row was last pulled from P6."),
              ]),
            T("p6_activity", "Live API", 132761,
              "Activity-level detail. Slippage shows here BEFORE it aggregates into project-level "
              "variance, which makes this the early-warning table.",
              [
                  ("p6_object_id", "bigint UQ", "P6 internal object ID."),
                  ("activity_id, name", "text", "Activity code and description."),
                  ("status, type", "text", "Completed / In Progress / Not Started, and the P6 activity type (task, milestone, level of effort)."),
                  ("start_date, finish_date", "timestamp", "Current dates."),
                  ("planned_start_date, planned_finish_date", "timestamp", "Planned dates."),
                  ("actual_start_date, actual_finish_date", "timestamp", "What actually happened."),
                  ("baseline_start_date, baseline_finish_date", "timestamp", "Per-activity baseline. The three-way comparison (planned / actual / baseline) is where slippage is detected."),
                  ("planned_duration, actual_duration, remaining_duration", "numeric", "Duration in days."),
                  ("percent_complete", "numeric", "Activity completion."),
                  ("total_float, is_critical", "numeric / bool", "Per-activity float and critical-path flag. A rising share of critical activities leads project SPI."),
                  ("wbs_object_id, wbs_code, wbs_name", "bigint / text", "Position in the work breakdown structure. The WBS code convention is what allows activity-level matching to SAP."),
                  ("project_object_id", "bigint FK", "Parent project -> p6_project.p6_object_id."),
                  ("last_synced_at", "timestamp", "Last pull from P6."),
              ]),
            T("p6_resource_assignment", "Live API", 300153,
              "Planned versus actual units per activity - progress measured by quantity rather than "
              "by duration. The two can diverge, and that divergence is meaningful.",
              [
                  ("p6_object_id", "bigint UQ", "P6 assignment object ID."),
                  ("activity_object_id", "bigint FK", "-> p6_activity.p6_object_id."),
                  ("project_object_id", "bigint FK", "-> p6_project.p6_object_id."),
                  ("resource_type", "text", "Labor, Nonlabor or Material."),
                  ("resource_name", "text", "Named resource."),
                  ("planned_units, actual_units", "numeric", "Planned and consumed units."),
                  ("last_synced_at", "timestamp", "Last pull from P6."),
              ]),
            T("p6_wbs_node", "Live API", 17078,
              "The work breakdown hierarchy. Nodes flagged as blocks correspond to physical site "
              "sub-areas, which is how schedule data lines up with SAP and transmission block references.",
              [
                  ("p6_object_id", "bigint UQ", "WBS node object ID."),
                  ("project_object_id", "bigint", "Owning project."),
                  ("wbs_code, wbs_name", "text", "Node code and label."),
                  ("parent_object_id", "bigint", "Parent node - the tree is walked through this."),
                  ("is_block, block_number", "bool / int", "Marks the node as a physical construction block and records its number."),
                  ("upload_time", "timestamp", "When the node was loaded."),
              ]),
            T("p6_baseline_project", "Live API", 197,
              "Full snapshots of the plan at each baselining event - one row per baseline, so drift "
              "can be measured across the life of the project rather than against a single stored plan.",
              [
                  ("p6_object_id", "bigint UQ", "Baseline object ID."),
                  ("original_project_object_id", "bigint FK", "-> p6_project.p6_object_id."),
                  ("baseline_type_name, name", "text", "Baseline type (for example 'Project Baseline') and its label."),
                  ("planned_start_date, start_date, finish_date, scheduled_finish_date", "timestamp", "Baseline schedule dates."),
                  ("planned_duration, actual_duration, remaining_duration", "numeric", "Baseline durations."),
                  ("planned_cost, actual_total_cost, remaining_total_cost, baseline_total_cost", "numeric", "Baseline cost position."),
                  ("activity_count and status counts", "int", "Activity totals at baseline time."),
                  ("current_budget, original_budget, status", "numeric / text", "Budget position and baseline state."),
                  ("last_synced_at", "timestamp", "Last pull from P6."),
              ]),
            T("p6_activity_risk", "Live API - EMPTY", 0,
              "The explicit risk register from P6 - risks authored by planners against specific "
              "activities. Currently unpopulated: the sync path exists but the source register has no "
              "entries yet. When it fills, it should be the first place a risk question looks, ahead "
              "of any float- or SPI-based inference.",
              [
                  ("activity_object_id, project_object_id", "bigint FK", "Activity and project the risk is attached to."),
                  ("risk_id, risk_name, risk_object_id", "text / bigint", "Risk register entry identity."),
                  ("activity_id, activity_name", "text", "Denormalised activity labels for direct display."),
                  ("last_synced_at", "timestamp", "Last pull from P6."),
              ]),
        ],
    },
    {
        "key": "Materials & cost (SAP)",
        "blurb": "Procurement, inventory, consumption and committed cost. Loaded from standard SAP "
                 "report extracts rather than a live connection, so freshness follows the extract "
                 "date, not the clock.",
        "source": "Excel extracts in Data/NEW31/ - MB52 (inventory), ME2J + ZPSPS007 (purchase orders), "
                  "MB51 (consumption), ZPSPS007 (actual/commitment) - scripts/ingest_sap_data.py, "
                  "scripts/ingest_slr_data.py",
        "tables": [
            T("mt_poamount", "ME2J / ZPSPS007", 69557,
              "Purchase order line items. The pending-delivery columns here answer most 'what is "
              "holding up material' questions.",
              [
                  ("purchasing_document", "text", "PO number."),
                  ("company_code, plant_code", "text", "SAP org units. [SAP KEY 1]"),
                  ("wbs_element", "text", "Cost object the PO is charged to. [SAP KEY 2]"),
                  ("material_code, material_name, short_text", "text", "Material identity and line description."),
                  ("material_type", "text", "SAP material classification."),
                  ("vendor_code, vendor_name, buyer_name", "text", "Supplier and the buyer accountable for the PO."),
                  ("order_quantity, po_quantities", "numeric", "Quantity ordered."),
                  ("mw_multiplication_factor, po_quantities_mw", "numeric", "Conversion factor and the order expressed in MW-equivalent, so it can be compared to schedule and capacity figures."),
                  ("net_order_value, net_order_value_inr, currency", "numeric / text", "PO value in document currency and in INR."),
                  ("delivered_qty, quantity_received", "numeric", "Quantity received to date."),
                  ("delivered_value_inr_cr", "numeric", "Delivered value in INR crore."),
                  ("still_to_deliver_qty, still_to_be_delivered_qty", "numeric", "Outstanding quantity - the core pending signal."),
                  ("still_to_deliver_inr", "numeric", "Value still outstanding."),
                  ("delivery_date, delivery_completed_flag", "timestamp / text", "Expected delivery and whether the line is closed. Outstanding quantity past this date is overdue."),
                  ("document_date", "timestamp", "PO creation date. Old open POs are a different risk pattern from new ones."),
                  ("deletion_indicator", "text", "Marks cancelled lines - exclude these from totals."),
                  ("storage_location, block_plot_name", "text", "Destination store and site block."),
                  ("upload_time", "timestamp", "When the extract was loaded."),
              ]),
            T("mt_materialdocument", "MB51", 21473,
              "Material movements - goods issued to site and their reversals.",
              [
                  ("material_document", "text", "SAP material document number."),
                  ("movement_type", "text", "Restricted to 221/222 (issue to WBS and its reversal) and 261/262 (issue to order and its reversal)."),
                  ("material_code, material_name, material_description", "text", "Material identity."),
                  ("plant_code, wbs_element", "text", "SAP join keys."),
                  ("quantity, base_unit", "numeric / text", "Volume moved and its unit."),
                  ("amount_in_lc, amount_in_lc_cr", "numeric", "Value in local currency, and the same in INR crore."),
                  ("posting_date", "timestamp", "When the movement was booked. Drives consumption-rate trends."),
                  ("storage_location, block_plot_name, purchase_order", "text", "Where it came from and which PO it traces to."),
                  ("upload_time", "timestamp", "Extract load timestamp."),
              ],
              callout="Netting rule: 222 and 262 are reversals of 221 and 261. Summing quantity across "
                      "all movement types double-counts corrected entries - net consumption is issues "
                      "minus reversals."),
            T("mt_slr_data", "ZPSPS007", 3991,
              "Actual spend and open commitment by WBS element - the cost-side counterpart to the PO table.",
              [
                  ("po_document", "text", "Accounting document / PO reference."),
                  ("description", "text", "Line description."),
                  ("vendor_name", "text", "Supplier."),
                  ("actual_amount", "numeric", "Cost actually booked."),
                  ("commitment_amount", "numeric", "Committed but not yet booked. Actual plus commitment is the true exposure."),
                  ("wbs_element", "text", "Cost object. [SAP KEY 2]"),
                  ("type", "text", "Document type - POrd (purchase order), PReq (requisition) and similar."),
                  ("plant_code", "text", "Derived from the first six characters of the WBS element. [SAP KEY 1]"),
                  ("upload_time", "timestamp", "Extract load timestamp."),
              ]),
            T("mt_inventory", "MB52", 3482,
              "Live stock on hand at site stores.",
              [
                  ("material_code, material_name, material_description", "text", "Material identity."),
                  ("plant_code, plant_name, wbs_element", "text", "SAP join keys and the readable plant name."),
                  ("unrestricted_qty, quantity_inv", "numeric", "Usable stock on hand."),
                  ("value_unrestricted", "numeric", "Value of that stock - working capital sitting at site."),
                  ("quantity_mw", "numeric", "Stock expressed in MW-equivalent, directly comparable to capacity and schedule figures."),
                  ("mw_multiplication_factor", "numeric", "Conversion factor used above."),
                  ("base_unit", "text", "Unit of measure."),
                  ("special_stock, material_type, material_group", "text", "SAP stock classification."),
                  ("storage_location_mapping", "text", "Physical store location."),
                  ("movement_type_validation", "text", "Validation flag carried from ingestion."),
                  ("vendor_code, purchase_order, posting_date", "text / timestamp", "Origin of the stock."),
                  ("upload_time", "timestamp", "Extract load timestamp."),
              ],
              callout="Filter caveat: ingestion loads only rows with stock above zero. A material with "
                      "no row here may be genuinely out of stock or may simply never have been stocked "
                      "- the table cannot distinguish the two, so 'zero' should not be reported from "
                      "absence alone."),
            T("mt_trialrun", "Derived from P6", 2303,
              "Commissioning and trial-run milestones, extracted from P6 activities into a block-level "
              "capacity register. Slipping trial-run dates are the clearest late-stage indicator for "
              "COD risk. Note the mt_ prefix is misleading - this is P6-derived, not SAP.",
              [
                  ("activity_id", "text UQ", "Source P6 activity."),
                  ("activity_name", "text", "Milestone description."),
                  ("project_name, project_name_p6, project_name_block", "text", "Project labels at project and block level."),
                  ("trial_run_start, trial_run_finish", "timestamp", "Commissioning window."),
                  ("tr_quantity_mw, unit_of_measure", "numeric / text", "MW commissioned in the trial run. Accumulates towards project nameplate capacity."),
                  ("spv_plant_code", "text", "SAP plant code. [SAP KEY 1]"),
                  ("is_start_before_upload", "text", "Data-quality flag - the trial run started before the data was loaded. Treat as a loading warning, not a real anomaly."),
                  ("upload_time", "timestamp", "Load timestamp."),
              ]),
            T("mt_requirement", "RETIRED", 0,
              "A planned material-demand register that has been retired - the table still exists in the "
              "database but has no rows and no model definition in the application. It should be "
              "dropped in the next migration."),
        ],
    },
    {
        "key": "Transmission network",
        "blurb": "Evacuation infrastructure - substations, line segments and their construction stage. "
                 "A generation project can be complete and still unable to export power, which is what "
                 "this domain makes visible.",
        "source": "Transmission portal export (tc_dump.json) - scripts/load_tc_data_from_json.py, "
                  "services/tc_sync.py",
        "tables": [
            T("tc_network_edge", "Portal JSON", 565,
              "Transmission line segments - the highest-value table in this domain. One delayed line "
              "can be the root cause behind several projects' delays at once.",
              [
                  ("edge_id", "text", "Line segment identifier."),
                  ("region", "text", "Khavda or Rajasthan."),
                  ("from_node, from_label, to_node, to_label", "text", "The two substations the line connects."),
                  ("contractor, voltage, length", "text", "Line attributes. Length is stored as text because the source is inconsistent."),
                  ("foundation, erection, stringing", "text", "Stage-gate construction progress. The furthest-behind stage is the real bottleneck on the line."),
                  ("status, normalized_status", "text", "Raw portal status and the cleaned version used for filtering."),
                  ("expected_date, scd, charged_date", "text", "Expected energisation, scheduled commissioning date, and when the line was actually charged."),
                  ("is_delayed", "bool", "Delay flag computed during ingestion - use it directly rather than re-deriving."),
                  ("projects", "text (JSON)", "All projects served by the line. A line commonly serves several."),
                  ("mapping_id", "int FK", "-> project_mapping.id. Representative link only - check the projects column for the full set. [TC KEY]"),
                  ("upload_time", "timestamp", "Load timestamp."),
              ]),
            T("tc_project_entry", "Portal JSON", 396,
              "Transmission scope by project and phase, at block granularity.",
              [
                  ("region, project, phase", "text", "Where and which phase of the scope."),
                  ("kps, pss", "text", "Pooling substation references."),
                  ("block, breakup", "text", "Site block and its scope breakdown."),
                  ("mw", "numeric", "Capacity carried on this entry."),
                  ("mapping_id", "int FK", "-> project_mapping.id. [TC KEY]"),
                  ("upload_time", "timestamp", "Load timestamp."),
              ]),
            T("tc_network_node", "Portal JSON", 56,
              "Substations and network nodes, including map coordinates for the network view.",
              [
                  ("node_id, label", "text", "Node identity and display name."),
                  ("region", "text", "Khavda or Rajasthan."),
                  ("type, status", "text", "Node classification and current state."),
                  ("x, y", "numeric", "Layout coordinates for rendering the network diagram."),
                  ("upload_time", "timestamp", "Load timestamp."),
              ]),
        ],
    },
    {
        "key": "Quality (Pulse)",
        "blurb": "Non-conformances and inspection requests. A project can look healthy on schedule and "
                 "cost while carrying critical quality exposure, so this domain is read alongside the "
                 "others rather than after them.",
        "source": "Pulse OData API - /pulse-api/Ncs and /pulse-api/Rfis - services/pulse_service.py",
        "tables": [
            T("pulse_rfi", "Live API", 43798,
              "Requests for inspection. Flattened from nested OData so dashboards can filter without "
              "traversing relations.",
              [
                  ("pulse_id", "text UQ", "Pulse record UUID - the deduplication key on sync."),
                  ("rfi_label", "text", "Human-readable RFI reference."),
                  ("status, status_label", "text", "Workflow state, coded and readable."),
                  ("current_handler", "text", "Whose queue it is sitting in."),
                  ("cluster_name, project_name, project_id, project_type", "text", "Location hierarchy from the portal."),
                  ("spv_name, worklocation_name, workarea_name", "text", "Entity, site and block."),
                  ("contractor_name, vendor_name, engineer_name, quality_name", "text", "People and firms attached to the inspection."),
                  ("package_name, inspection_point_name", "text", "Work package and the specific inspection point."),
                  ("created_at, updated_at", "timestamp", "Portal-side timestamps. Ageing is measured from these."),
                  ("last_synced_at", "timestamp", "Last pull from Pulse."),
              ]),
            T("pulse_nc", "Live API", 576,
              "Non-conformance records, including the financial penalty raised against the contractor.",
              [
                  ("pulse_id", "text UQ", "Pulse record UUID."),
                  ("nc_label", "text", "NC reference number."),
                  ("status, status_label", "text", "Raised, in review, approved, rejected or completed."),
                  ("category", "text", "Critical or Non Critical - the primary severity split."),
                  ("defect_type, description, quantity", "text / numeric", "What the defect is and how much is affected."),
                  ("debit, debit_reason", "numeric / text", "Financial penalty raised, and its justification."),
                  ("current_handler", "text", "Contractor, execution engineer or quality inspector."),
                  ("ad_hoc, archived, version", "bool / int", "Record flags and revision number."),
                  ("cluster_name, project_name, project_id, project_type", "text", "Location hierarchy."),
                  ("spv_name, worklocation_name, workarea_name", "text", "Entity, site and block."),
                  ("contractor_name, vendor_name, vendor_code, engineer_name, quality_name", "text", "Accountability - who raised it and who owes the fix."),
                  ("package_name, subpackage_name, activity_name, subactivity_name", "text", "Work breakdown, flattened four levels deep from the portal hierarchy."),
                  ("service_order_number", "text", "Linked service order."),
                  ("created_at, updated_at, approved_at", "timestamp", "Lifecycle timestamps. Time-to-approval is derived from these."),
                  ("last_synced_at", "timestamp", "Last pull from Pulse."),
              ]),
        ],
    },
    {
        "key": "Invoicing",
        "blurb": "Contractor invoices and their approval position, joined back to projects through the "
                 "purchase order.",
        "source": "E-Invoice portal REST API (OAuth) - scripts/sync_einvoice_live.py - PO lookup built "
                  "from ME2J and ZPSPS007",
        "tables": [
            T("einvoice_records", "Live API", 362,
              "One row per invoice, with the workflow fields needed to see what is stuck and with whom.",
              [
                  ("invoiceNo, invoiceCode, invoiceRequestID", "text / int", "Invoice identity in the portal."),
                  ("vendorName, sapVendorCode", "text", "Supplier, and the code that ties them back to SAP."),
                  ("projectType, packageName", "text", "Asset type and work package billed."),
                  ("workLocation, site", "text", "Where the work was performed."),
                  ("p6ProjectName", "text", "Project name as matched against P6."),
                  ("workOrderNo", "text", "Purchase order number - the join back to SAP and to the project. [JOIN]"),
                  ("workDescription", "text", "What was billed."),
                  ("invoiceAmount, soAmount", "numeric", "Invoice value and the service-order value it draws against."),
                  ("statusDesc, stage", "text", "Approval status and current workflow stage."),
                  ("isPending", "bool", "Whether the invoice is still awaiting action."),
                  ("currentApprover, latestAction", "text", "Who holds it now and what happened last."),
                  ("invoiceDate, submittedOn, createdAt, completionDate", "timestamp", "Lifecycle dates. Approval ageing is derived from submission to completion."),
              ]),
            T("mt_einvoice_po_lookup", "Derived index", 11268,
              "A dedicated PO-to-WBS index built from both ME2J and ZPSPS007. It exists so that POs "
              "present in the invoice feed but missing from the main SAP tables can still be resolved "
              "to a project WITHOUT polluting the materials dashboards with phantom rows.",
              [
                  ("purchasing_document", "varchar(50)", "PO number - matches einvoice_records.workOrderNo."),
                  ("wbs_element", "varchar(255)", "WBS element the PO is charged to. Its first six characters give the plant."),
              ]),
        ],
    },
    {
        "key": "Platform tables",
        "blurb": "Generated by the application itself rather than ingested from a source system: users, "
                 "alerts, the assistant's conversation history, and a computed-metrics cache.",
        "source": "Akasha application runtime",
        "tables": [
            T("notification", "Application", 4148,
              "Change alerts raised when synced data moves - date slips, budget breaches, status changes.",
              [
                  ("project_name, module", "text", "Which project, and which domain raised it (P6 or Transmission)."),
                  ("change_type", "text", "Date change, budget exceeded, status update or critical date slip."),
                  ("category", "text", "Budgets, COD, Trials or Dates."),
                  ("message", "text", "Alert text shown to the user."),
                  ("block, activity_name", "text", "Where in the project the change occurred."),
                  ("old_value, new_value, reason", "text", "What changed, and why if known."),
                  ("action_status", "text", "Pending, Acknowledged or Resolved."),
                  ("p6_object_id, p6_type", "bigint / text", "Link back to the P6 record, so an acknowledged change can be pushed upstream."),
                  ("ai_suggestion", "text", "Pre-generated recommended action."),
                  ("is_read, created_at", "bool / timestamp", "Read state and when it was raised."),
              ]),
            T("notification_thread", "Application - EMPTY", 0,
              "Discussion replies against a notification - notification_id, sender, message, created_at. "
              "Deleted with the parent notification. Not yet in use."),
            T("chat_message", "Application", 194,
              "Assistant conversation turns, each stamped with the data it drew on - so any answer can "
              "be traced back to its sources and their freshness.",
              [
                  ("session_id", "text FK", "-> chat_session.session_id."),
                  ("role, content", "text", "User or assistant, and the message body."),
                  ("intent_type", "text", "Factual, analytical, advisory or document."),
                  ("project_ids, data_domains", "text", "Which projects and which domains (P6, SAP, TC) the answer touched."),
                  ("data_as_of", "timestamp", "Freshness of the underlying data at answer time."),
                  ("sources_used", "jsonb", "Tables read and whether the metrics cache was hit."),
                  ("latency_ms, created_at", "int / timestamp", "Response time and when it was produced."),
              ]),
            T("chat_session", "Application", 32,
              "Server-side conversations: session_id, title (generated from the first message), "
              "is_active, created_at, updated_at. Messages cascade-delete with the session."),
            T("metrics_cache", "Application", 4,
              "Computed dashboard, project-360 and variance results held per project so they are not "
              "recomputed on every question. Each entry stores the P6, SAP and transmission sync "
              "timestamps it was built from, so it can be invalidated the moment upstream data actually "
              "moves - project_id, cache_key, data, computed_at, p6_synced_at, sap_synced_at, tc_synced_at."),
            T("chat_feedback", "Application - EMPTY", 0,
              "Thumbs up/down and written corrections against an answer - message_id, feedback_type, "
              "correction_text, project_id, question_pattern. Corrections are intended to feed back into "
              "later answers on the same project and question shape."),
            T("akasha_user", "Application - EMPTY", 0,
              "Application accounts - username, password_hash, display_name, email, is_active, "
              "created_at, and role (executive, pmag, projects, tc_ordering, tc_stores). Empty in this "
              "environment; users are seeded per deployment."),
            T("alembic_version", "Migrations", 1,
              "Schema migration bookkeeping. Not application data."),
        ],
    },
]

CONVENTIONS = [
    ("Freshness differs by domain",
     "P6, Pulse and E-Invoice sync over live APIs and carry a last_synced_at per row. SAP and "
     "transmission data arrive as file extracts and carry upload_time - they are only as current as "
     "the last extract, which is a manual step. P6 additionally has data_date, the schedule's own "
     "cutoff: progress recorded after that date is not reflected even if the sync ran minutes ago."),
    ("Naming",
     "Tables prefixed mt_ come from SAP or SAP-adjacent material tracking, p6_ from Primavera, tc_ "
     "from the transmission portal, pulse_ from the quality system. einvoice_records keeps the "
     "portal's own camelCase column names rather than being renamed on ingest."),
    ("Currency and units",
     "Amounts are INR unless a column name says otherwise; columns suffixed _cr are in crore. Capacity "
     "is MW throughout - _mwac and _mwdc distinguish AC and DC ratings, and SAP quantities carry a "
     "matching _mw conversion so material can be compared against schedule and capacity on the same scale."),
    ("Empty tables",
     "Five tables currently hold no rows: p6_activity_risk, notification_thread, chat_feedback and "
     "akasha_user are wired up but not yet populated in this environment. mt_requirement is retired - "
     "it has no model definition left in the application and should be dropped."),
    ("Derived values are not stored",
     "Percent complete against capacity, delivery ageing, net consumption, stage-gate progress and "
     "schedule slack are calculated at query time from the columns above. They are not columns in the "
     "database, so a figure on a dashboard will not always be findable as a stored field."),
]


# ==========================================================
# Excel
# ==========================================================

THIN = Side(style="thin", color=LINE)
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def style_header(ws, row, ncols, fill=COPPER):
    for c in range(1, ncols + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = Font(bold=True, color="FFFFFF", size=10)
        cell.fill = PatternFill("solid", fgColor=fill)
        cell.alignment = Alignment(vertical="center", horizontal="left")
        cell.border = BORDER
    ws.row_dimensions[row].height = 22


def set_widths(ws, widths):
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w


def write_rows(ws, start, rows, wrap_col=None):
    r = start
    for row in rows:
        for i, val in enumerate(row, start=1):
            cell = ws.cell(row=r, column=i, value=val)
            cell.border = BORDER
            cell.alignment = Alignment(vertical="top",
                                       wrap_text=(wrap_col is not None and i >= wrap_col))
            if isinstance(val, int):
                cell.number_format = "#,##0"
        r += 1
    return r


def build_excel(path):
    wb = Workbook()

    # ---------- Overview ----------
    ws = wb.active
    ws.title = "Overview"
    ws.sheet_view.showGridLines = False
    set_widths(ws, [26, 22, 14, 84])

    ws["A1"] = "Akasha Platform - Data Schema Reference"
    ws["A1"].font = Font(bold=True, size=16, color=INK)
    ws["A2"] = (f"Every table the platform currently stores, what feeds it, and how the domains join "
                f"together. Generated {GENERATED} from the live PostgreSQL database 'Akasha' and "
                f"backend/models.py.")
    ws["A2"].font = Font(size=10, color="61717E")
    ws.merge_cells("A2:D2")
    ws.row_dimensions[2].height = 30
    ws["A2"].alignment = Alignment(wrap_text=True, vertical="top")

    r = 4
    ws.cell(row=r, column=1, value="AT A GLANCE").font = Font(bold=True, size=10, color=TEAL)
    r += 1
    for label, val in SUMMARY:
        ws.cell(row=r, column=1, value=label).font = Font(size=10, color="61717E")
        c = ws.cell(row=r, column=2, value=val)
        c.font = Font(bold=True, size=11, color=INK)
        r += 1

    r += 1
    ws.cell(row=r, column=1, value="ALL TABLES").font = Font(bold=True, size=10, color=TEAL)
    r += 1
    hdr = r
    for i, h in enumerate(["Domain", "Table", "Rows", "Purpose"], start=1):
        ws.cell(row=hdr, column=i, value=h)
    style_header(ws, hdr, 4)
    r += 1

    overview_rows = []
    for d in DOMAINS:
        for t in d["tables"]:
            overview_rows.append((d["key"], t["name"], t["rows"], t["note"]))
    r = write_rows(ws, r, overview_rows, wrap_col=4)
    ws.freeze_panes = ws.cell(row=hdr + 1, column=1)
    ws.auto_filter.ref = f"A{hdr}:D{r - 1}"

    # ---------- Join Keys ----------
    ws = wb.create_sheet("Join Keys")
    ws.sheet_view.showGridLines = False
    set_widths(ws, [28, 34, 86])
    ws["A1"] = "How the domains join"
    ws["A1"].font = Font(bold=True, size=14, color=INK)
    ws["A2"] = ("Three source systems name the same physical project three different ways. "
                "project_mapping is the resolver that reconciles them - nothing should be matched "
                "across domains by project name.")
    ws["A2"].font = Font(size=10, color="61717E")
    ws.merge_cells("A2:C2")
    ws["A2"].alignment = Alignment(wrap_text=True, vertical="top")
    ws.row_dimensions[2].height = 30

    hdr = 4
    for i, h in enumerate(["Key on project_mapping", "Resolves to", "Applies to"], start=1):
        ws.cell(row=hdr, column=i, value=h)
    style_header(ws, hdr, 3)
    r = write_rows(ws, hdr + 1, JOIN_KEYS, wrap_col=3)

    r += 1
    ws.cell(row=r, column=1, value=RETRIEVAL_RULE).font = Font(bold=True, size=10, color=COPPER)
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=3)
    ws.cell(row=r, column=1).alignment = Alignment(wrap_text=True, vertical="top")
    ws.row_dimensions[r].height = 44

    # ---------- One sheet per domain ----------
    sheet_names = {
        "Project master": "Project Master",
        "Schedule & progress (P6)": "P6 Schedule",
        "Materials & cost (SAP)": "SAP Materials",
        "Transmission network": "Transmission",
        "Quality (Pulse)": "Quality",
        "Invoicing": "Invoicing",
        "Platform tables": "Platform",
    }

    for d in DOMAINS:
        ws = wb.create_sheet(sheet_names[d["key"]])
        ws.sheet_view.showGridLines = False
        set_widths(ws, [46, 24, 96])

        ws["A1"] = d["key"]
        ws["A1"].font = Font(bold=True, size=14, color=INK)
        ws["A2"] = d["blurb"]
        ws["A2"].font = Font(size=10, color="61717E")
        ws.merge_cells("A2:C2")
        ws["A2"].alignment = Alignment(wrap_text=True, vertical="top")
        ws.row_dimensions[2].height = 30
        ws["A3"] = "SOURCE:  " + d["source"]
        ws["A3"].font = Font(size=9, color=TEAL)
        ws.merge_cells("A3:C3")
        ws["A3"].alignment = Alignment(wrap_text=True, vertical="top")
        ws.row_dimensions[3].height = 26

        r = 5
        for t in d["tables"]:
            label = f"{t['name']}    -    {t['tag']}    -    {t['rows']:,} rows"
            cell = ws.cell(row=r, column=1, value=label)
            cell.font = Font(bold=True, size=11, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor=INK)
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=3)
            ws.row_dimensions[r].height = 20
            r += 1

            cell = ws.cell(row=r, column=1, value=t["note"])
            cell.font = Font(size=10, italic=True, color="33424E")
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=3)
            ws.row_dimensions[r].height = max(16, 14 * (len(t["note"]) // 110 + 1))
            r += 1

            if t["cols"]:
                for i, h in enumerate(["Column", "Type", "Meaning"], start=1):
                    ws.cell(row=r, column=i, value=h)
                style_header(ws, r, 3, fill=TEAL)
                r += 1
                r = write_rows(ws, r, t["cols"], wrap_col=3)

            if t["callout"]:
                cell = ws.cell(row=r, column=1, value=t["callout"])
                cell.font = Font(size=10, bold=True, color=COPPER)
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=3)
                ws.row_dimensions[r].height = 40
                r += 1

            r += 1

    # ---------- All Columns (flat, filterable) ----------
    ws = wb.create_sheet("All Columns")
    ws.sheet_view.showGridLines = False
    set_widths(ws, [24, 26, 46, 24, 92])
    for i, h in enumerate(["Domain", "Table", "Column", "Type", "Meaning"], start=1):
        ws.cell(row=1, column=i, value=h)
    style_header(ws, 1, 5)
    flat = []
    for d in DOMAINS:
        for t in d["tables"]:
            for c in t["cols"]:
                flat.append((d["key"], t["name"], c[0], c[1], c[2]))
    r = write_rows(ws, 2, flat, wrap_col=5)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:E{r - 1}"

    # ---------- Conventions ----------
    ws = wb.create_sheet("Conventions")
    ws.sheet_view.showGridLines = False
    set_widths(ws, [34, 110])
    ws["A1"] = "Conventions & caveats"
    ws["A1"].font = Font(bold=True, size=14, color=INK)
    ws["A2"] = "What someone reading a number out of this database needs to know before quoting it."
    ws["A2"].font = Font(size=10, color="61717E")

    hdr = 4
    for i, h in enumerate(["Topic", "Detail"], start=1):
        ws.cell(row=hdr, column=i, value=h)
    style_header(ws, hdr, 2)
    write_rows(ws, hdr + 1, CONVENTIONS, wrap_col=2)

    wb.save(path)
    return path


# ==========================================================
# Word
# ==========================================================

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_docx(path):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    for lvl, size in ((0, 22), (1, 15), (2, 12)):
        st = doc.styles[f"Heading {lvl + 1}"] if lvl else doc.styles["Title"]
        st.font.color.rgb = RGBColor.from_string(INK)
        st.font.size = Pt(size)
        st.font.name = "Calibri"

    doc.add_paragraph("AKASHA PLATFORM  ·  DATA LAYER").runs[0].font.color.rgb = \
        RGBColor.from_string(COPPER)
    doc.paragraphs[-1].runs[0].font.size = Pt(9)
    doc.paragraphs[-1].runs[0].font.bold = True

    doc.add_heading("Data Schema Reference", 0)

    p = doc.add_paragraph(
        "Every table the platform currently stores, what feeds it, and how the domains join "
        "together. Generated from the live PostgreSQL database 'Akasha' and the SQLAlchemy models "
        "in backend/models.py.")
    p.runs[0].font.size = Pt(11)

    meta = doc.add_paragraph()
    run = meta.add_run(
        f"Database: PostgreSQL · Akasha     |     Tables: 27     |     "
        f"As of: {GENERATED}     |     Source of truth: backend/models.py")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("61717E")

    # at a glance
    doc.add_heading("At a glance", 2)
    tbl = doc.add_table(rows=1, cols=len(SUMMARY))
    tbl.style = "Table Grid"
    for i, (label, val) in enumerate(SUMMARY):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        pv = cell.paragraphs[0]
        rv = pv.add_run(val)
        rv.bold = True
        rv.font.size = Pt(15)
        pl = cell.add_paragraph()
        rl = pl.add_run(label.upper())
        rl.font.size = Pt(7.5)
        rl.font.color.rgb = RGBColor.from_string("61717E")

    # join keys
    doc.add_heading("How the domains join", 1)
    doc.add_paragraph(
        "Three source systems name the same physical project three different ways. project_mapping "
        "is the resolver that reconciles them - nothing should be matched across domains by project "
        "name.")
    add_table(doc, ["Key on project_mapping", "Resolves to", "Applies to"], JOIN_KEYS,
              widths=[1.6, 1.9, 3.4])
    rule = doc.add_paragraph()
    rr = rule.add_run(RETRIEVAL_RULE)
    rr.bold = True
    rr.font.size = Pt(9)
    rr.font.color.rgb = RGBColor.from_string(COPPER)

    # domains
    for d in DOMAINS:
        doc.add_page_break()
        doc.add_heading(d["key"], 1)
        doc.add_paragraph(d["blurb"])
        sp = doc.add_paragraph()
        sr = sp.add_run("SOURCE:  " + d["source"])
        sr.font.size = Pt(8)
        sr.font.color.rgb = RGBColor.from_string(TEAL)

        for t in d["tables"]:
            doc.add_heading(t["name"], 2)
            tp = doc.add_paragraph()
            tr = tp.add_run(f"{t['tag']}  ·  {t['rows']:,} rows")
            tr.font.size = Pt(8)
            tr.bold = True
            tr.font.color.rgb = RGBColor.from_string("61717E")

            np_ = doc.add_paragraph(t["note"])
            np_.runs[0].italic = True
            np_.runs[0].font.size = Pt(9.5)

            if t["cols"]:
                add_table(doc, ["Column", "Type", "Meaning"], t["cols"],
                          widths=[2.0, 1.1, 3.8], header_fill=TEAL)

            if t["callout"]:
                cp = doc.add_paragraph()
                cr = cp.add_run(t["callout"])
                cr.bold = True
                cr.font.size = Pt(9)
                cr.font.color.rgb = RGBColor.from_string(COPPER)

    # conventions
    doc.add_page_break()
    doc.add_heading("Conventions & caveats", 1)
    doc.add_paragraph(
        "What someone reading a number out of this database needs to know before quoting it.")
    for topic, detail in CONVENTIONS:
        doc.add_heading(topic, 2)
        doc.add_paragraph(detail)

    f = doc.add_paragraph()
    fr = f.add_run(
        f"Akasha Platform data schema · generated {GENERATED} from the live PostgreSQL database and "
        f"backend/models.py. Row counts are a point-in-time snapshot and change with every sync. "
        f"Regenerate this document whenever a model or ingestion script changes - it will otherwise "
        f"go stale silently.")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string("61717E")

    doc.save(path)
    return path


def add_table(doc, headers, rows, widths=None, header_fill=COPPER):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False

    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h.upper())
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        shade(cell, header_fill)

    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
            if i == 0:
                run.font.name = "Consolas"
                run.bold = True

    if widths:
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return tbl


if __name__ == "__main__":
    xlsx = build_excel(os.path.join(ROOT, "Akasha_Data_Schema.xlsx"))
    docx_path = build_docx(os.path.join(ROOT, "Akasha_Data_Schema.docx"))
    print("Wrote:")
    print("  " + xlsx)
    print("  " + docx_path)
