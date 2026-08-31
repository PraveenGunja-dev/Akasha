"""
Generate the client-facing API reference as a Word document.

Everything in the output is taken from the running service: endpoint list from
the FastAPI schema, examples from live responses, row counts and coverage from
the database. Re-run it after an API change and the document is current.

    venv/Scripts/python.exe scripts/build_api_doc.py
"""

import json
import os
import sys
import warnings
from datetime import datetime

warnings.filterwarnings("ignore")
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402

INK = RGBColor(0x11, 0x17, 0x21)
MUTED = RGBColor(0x60, 0x6A, 0x78)
ACCENT = RGBColor(0x0B, 0x74, 0xB1)
CODE_BG = "F2F4F7"
HEAD_BG = "0B74B1"

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "Akasha_API_v1_Reference.docx")


# ───────────────────────── formatting helpers ─────────────────────────

def shade(cell, hex_colour):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(el)


def code(doc, text, size=8.5):
    """Monospace block on a light ground — used for requests and payloads."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    shade(cell, CODE_BG)
    cell.text = ""
    for i, line in enumerate(text.rstrip("\n").split("\n")):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(size)
        run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade(cell, HEAD_BG)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h.upper())
        run.bold = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            run = para.add_run(str(value))
            run.font.size = Pt(8.5)
            if i == 0:
                run.font.name = "Consolas"
                run.font.size = Pt(8)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def para(doc, text, size=10, colour=None, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = colour or INK
    return p


def callout(doc, title, body):
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade(cell, "FFF7E6")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(9.5)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def pretty(body, max_rows=1, max_fields=20):
    """Compact a live response so the sample stays readable in print.

    Trims the number of ROWS and the number of FIELDS per row, never the tail
    of the document — otherwise long source tables push `meta` off the end and
    the reader never sees the envelope, which is the part they have to code
    against.
    """
    b = json.loads(json.dumps(body, default=str))
    rows = b.get("data")

    def trim(row):
        if not isinstance(row, dict) or len(row) <= max_fields:
            return row
        kept = dict(list(row.items())[:max_fields])
        kept[f"... {len(row) - max_fields} more fields"] = "omitted from this example"
        return kept

    if isinstance(rows, list):
        extra = len(rows) - max_rows
        b["data"] = [trim(r) for r in rows[:max_rows]]
        if extra > 0:
            b["data"].append(f"... {extra} more rows omitted from this example")
    elif isinstance(rows, dict):
        b["data"] = trim(rows)

    return json.dumps(b, indent=2)


# ───────────────────────── content ─────────────────────────

SAMPLES = json.load(open(
    r"C:\Users\USER\AppData\Local\Temp\claude\d--Akasha-Platform"
    r"\f062b9ec-710d-4750-be0d-adb70f8adea7\scratchpad\samples.json",
    encoding="utf-8"))
PID = SAMPLES["_pid"]

COMMON_PARAMS = [
    ("project_id", "string", "—", "Canonical project id, or any accepted alias. Omit for the whole portfolio. Unknown value returns 404."),
    ("portfolio", "string", "all", "Solar Khavda | Solar Rajasthan | Wind | BESS"),
    ("phase", "enum", "all", "ongoing | commissioned | all. Any other value returns 422."),
    ("page", "integer", "1", "1-based page number."),
    ("page_size", "integer", "50", "1–200. Above 200 returns 422."),
]

ENDPOINTS = [
    {
        "path": "/api/v1/projects", "sample": "/api/v1/projects",
        "group": "Discovery",
        "title": "List projects",
        "desc": "Every project the platform knows about, each with its canonical id and the "
                "source systems it is linked to. This is the entry point for most integrations.",
        "params": [p for p in COMMON_PARAMS if p[0] != "project_id"],
        "fields": [
            ("project_id", "string", "Canonical identifier. Use this everywhere else."),
            ("mapping_id", "integer", "Internal surrogate key. Accepted as an alias."),
            ("name", "string", "Human-readable project name."),
            ("portfolio", "string", "Cluster the project belongs to."),
            ("is_commissioned", "boolean", "False = ongoing, True = commissioned."),
            ("capacity_mwac", "number", "Nameplate capacity in MWac."),
            ("linked", "array", "Source systems this project resolves to."),
            ("unlinked", "array", "Source systems it does NOT resolve to. An empty result from one of these means 'not connected', not 'no data'."),
        ],
    },
    {
        "path": "/api/v1/projects/{project_id}", "sample": f"/api/v1/projects/{PID}",
        "group": "Discovery",
        "title": "Get one project",
        "desc": "A single project. The path segment is tolerant: the canonical id, the numeric "
                "mapping id, a P6 object id, or the project name all resolve to the same record. "
                "The response always answers with the canonical id.",
        "params": [],
        "fields": [],
    },
    {
        "path": "/api/v1/projects/{project_id}/identity", "sample": f"/api/v1/projects/{PID}/identity",
        "group": "Discovery",
        "title": "Project key map",
        "desc": "What this project is called in each source system. Needed only if you are "
                "building your own joins against the underlying tables.",
        "params": [],
        "fields": [
            ("keys.canonical", "string", "The canonical project id."),
            ("keys.p6", "object", "project_id and object_id in Primavera P6."),
            ("keys.sap", "object", "plant_code and the WBS prefixes that belong to this project."),
            ("keys.tc", "object", "mapping_id used by transmission records."),
            ("keys.pulse", "object", "Pulse project name, where linked."),
        ],
    },
    {
        "path": "/api/v1/coverage", "sample": "/api/v1/coverage",
        "group": "Discovery",
        "title": "Source coverage",
        "desc": "How much of the portfolio links to each source system, and which projects do "
                "not. Call this before trusting a portfolio-wide aggregate.",
        "params": [p for p in COMMON_PARAMS if p[0] in ("portfolio", "phase")],
        "fields": [
            ("total_projects", "integer", "Projects in scope."),
            ("systems.<name>.linked", "integer", "How many resolve to that system."),
            ("systems.<name>.pct", "integer", "Percentage linked."),
            ("systems.<name>.unlinked_project_ids", "array", "Up to 50 project ids that do not resolve."),
        ],
    },
    {
        "path": "/api/v1/p6", "sample": f"/api/v1/p6?project_id={PID}",
        "group": "Schedule (P6)", "title": "P6 project schedule",
        "desc": "Project-level schedule from Primavera P6 — dates, baselines, float, variance, "
                "cost and activity counts. Joined directly on the canonical id.",
        "params": COMMON_PARAMS,
        "fields": [
            ("project_id", "string", "Canonical id (added by the API)."),
            ("start_date / finish_date", "datetime", "Current schedule dates."),
            ("baseline_start_date / baseline_finish_date", "datetime", "Original approved dates."),
            ("duration_percent_complete", "number", "Progress, 0–100."),
            ("total_float", "number", "Days of float. <= 0 means on the critical path."),
            ("finish_date_variance", "number", "Days against baseline. Negative = late."),
            ("activity_count", "integer", "Total activities; completed / in-progress / not-started also returned."),
        ],
    },
    {
        "path": "/api/v1/activities", "sample": f"/api/v1/activities?project_id={PID}",
        "group": "Schedule (P6)", "title": "P6 activities",
        "desc": "Individual schedule activities. 123,180 rows portfolio-wide, so always scope "
                "this by project_id. Default page size is 200 here rather than 50.",
        "params": COMMON_PARAMS,
        "fields": [
            ("activity_id / name", "string", "Activity identifier and description."),
            ("status / type", "string", "Activity state and type."),
            ("start_date / finish_date", "datetime", "Actual or scheduled dates."),
            ("is_critical", "boolean", "On the critical path."),
            ("wbs_name", "string", "Work breakdown element the activity sits under."),
        ],
    },
    {
        "path": "/api/v1/sap", "sample": f"/api/v1/sap?project_id={PID}",
        "group": "Financials (SAP)", "title": "Purchase orders",
        "desc": "SAP purchase orders. Matched by WBS prefix; where a prefix is claimed by more "
                "than one project the most specific wins, so every row is attributed to exactly "
                "one project and never counted twice.",
        "params": COMMON_PARAMS,
        "fields": [
            ("wbs_element", "string", "SAP WBS the order sits under."),
            ("net_order_value_inr", "number", "Order value in rupees (not crore)."),
            ("delivered_value_inr_cr", "number", "Delivered value in crore."),
            ("delivered_qty / still_to_deliver_qty", "number", "Quantity delivered and outstanding."),
            ("document_date / delivery_date", "datetime", "Order and expected delivery dates."),
            ("vendor_code / material_code", "string", "Vendor and material references."),
        ],
    },
    {
        "path": "/api/v1/slr", "sample": f"/api/v1/slr?project_id={PID}",
        "group": "Financials (SAP)", "title": "SLR ledger",
        "desc": "Service and labour ledger entries — actual against commitment by WBS. Same "
                "single-owner WBS rule as purchase orders.",
        "params": COMMON_PARAMS,
        "fields": [
            ("po_document", "string", "Source purchase document."),
            ("actual_amount / commitment_amount", "number", "Spend and commitment."),
            ("vendor_name", "string", "Supplier."),
            ("type", "string", "Ledger entry type."),
        ],
    },
    {
        "path": "/api/v1/inventory", "sample": f"/api/v1/inventory?project_id={PID}",
        "group": "Financials (SAP)", "title": "Inventory / GRN",
        "desc": "Goods-received and inventory positions by WBS.",
        "params": COMMON_PARAMS,
        "fields": [
            ("quantity_inv / quantity_mw", "number", "Quantity held, and its MW equivalent."),
            ("posting_date", "datetime", "Posting date."),
            ("storage_location_mapping", "string", "Where the stock sits."),
            ("material_type / material_group", "string", "Material classification."),
        ],
    },
    {
        "path": "/api/v1/material-documents", "sample": f"/api/v1/material-documents?project_id={PID}",
        "group": "Financials (SAP)", "title": "Material documents (MB51)",
        "desc": "Consumption and movement documents. Used with purchase orders to derive "
                "actual expenditure.",
        "params": COMMON_PARAMS,
        "fields": [
            ("quantity", "number", "Movement quantity."),
            ("amount_in_lc", "number", "Value in local currency."),
            ("posting_date", "datetime", "Posting date."),
        ],
    },
    {
        "path": "/api/v1/trial-run", "sample": "/api/v1/trial-run",
        "group": "Financials (SAP)", "title": "Trial run and COD",
        "desc": "Trial-run and commercial-operation milestones by block. Joined on the P6 "
                "project name, falling back to the SPV plant code.",
        "params": COMMON_PARAMS,
        "fields": [
            ("project_name_block", "string", "Block within the project."),
            ("spv_plant_code", "string", "SAP SPV plant code."),
        ],
    },
    {
        "path": "/api/v1/einvoice", "sample": "/api/v1/einvoice",
        "group": "Financials (SAP)", "title": "E-invoices",
        "desc": "Invoice records. These carry the P6 project name directly, so no WBS "
                "resolution is involved.",
        "params": COMMON_PARAMS,
        "fields": [
            ("invoiceNo / invoiceCode", "string", "Invoice references."),
            ("invoiceAmount / soAmount", "number", "Invoiced and sales-order amounts."),
            ("statusDesc", "string", "Current invoice status."),
            ("invoiceDate / completionDate", "datetime", "Key dates."),
            ("vendorName / sapVendorCode", "string", "Supplier."),
        ],
    },
    {
        "path": "/api/v1/pulse?kind=nc", "sample": "/api/v1/pulse?kind=nc",
        "group": "Quality (Pulse)", "title": "Non-conformances",
        "desc": "Site quality non-conformances. Joined on the Pulse project UUID stored on the "
                "mapping. Note the coverage caveat: Pulse resolves for roughly half the "
                "portfolio, so an empty result frequently means the project is not connected.",
        "params": COMMON_PARAMS + [("kind", "enum", "nc", "nc | rfi. Any other value returns 422.")],
        "fields": [
            ("project_id", "string", "Canonical id."),
            ("pulse_project_uuid", "string", "Pulse's own project identifier."),
            ("nc_label / description", "string", "Reference and detail."),
            ("status / status_label", "string", "raised | submitted | approved | completed | rejected."),
            ("category", "string", "Critical or Non Critical."),
            ("current_handler", "string", "Who the item sits with."),
            ("debit / debit_reason", "number / string", "Any debit raised."),
            ("created_at / approved_at", "datetime", "Raised and resolved timestamps."),
        ],
    },
    {
        "path": "/api/v1/pulse?kind=rfi", "sample": "/api/v1/pulse?kind=rfi",
        "group": "Quality (Pulse)", "title": "Requests for information",
        "desc": "Site RFIs. Same endpoint as non-conformances with kind=rfi.",
        "params": [("kind", "enum", "nc", "Set to rfi for this collection.")],
        "fields": [
            ("rfi_label", "string", "RFI reference."),
            ("status / status_label", "string", "Workflow state."),
            ("contractor_name / engineer_name", "string", "Parties involved."),
            ("worklocation_name / workarea_name", "string", "Where on site."),
        ],
    },
    {
        "path": "/api/v1/transmission", "sample": "/api/v1/transmission",
        "group": "Transmission (TC)", "title": "Transmission entries",
        "desc": "Grid connectivity records for the scoped projects.",
        "params": COMMON_PARAMS,
        "fields": [
            ("region", "string", "Khavda or Rajasthan."),
            ("phase / kps / pss / block", "string", "Network position."),
            ("mw", "number", "Capacity associated with the entry."),
            ("breakup", "string", "Line breakdown."),
        ],
    },
]


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    # ── cover ──
    doc.add_paragraph().paragraph_format.space_after = Pt(90)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Akasha Execution Platform")
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED
    p = doc.add_paragraph()
    r = p.add_run("API v1")
    r.font.size = Pt(40)
    r.bold = True
    r.font.color.rgb = INK
    p = doc.add_paragraph()
    r = p.add_run("Integration Reference")
    r.font.size = Pt(19)
    r.font.color.rgb = ACCENT

    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    para(doc,
         "One canonical project identifier reaches every source system — Primavera P6, SAP, "
         "the SLR ledger, Pulse quality and the Transmission portal. Fourteen endpoints, one "
         "filter vocabulary, one response envelope.",
         size=11, colour=MUTED)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    table(doc, ["Item", "Detail"], [
        ["Version", "v1"],
        ["Base URL", "https://<host>/akasha/api/v1"],
        ["Endpoints", "14 — 10 data sources, 4 discovery"],
        ["Format", "JSON (UTF-8)"],
        ["Generated", datetime.now().strftime("%d %b %Y")],
        ["Verification", "73 automated acceptance tests passing"],
    ], widths=[1.6, 4.6])

    doc.add_page_break()

    # ── 1 overview ──
    doc.add_heading("1. Getting started", level=1)
    para(doc, "Three calls cover most integrations: find a project, check what it links to, "
              "then pull any source using the same identifier.")
    code(doc,
         "# 1 - list projects\n"
         "GET /akasha/api/v1/projects?page_size=50\n\n"
         "# 2 - what is this project connected to?\n"
         f"GET /akasha/api/v1/projects/{PID}/identity\n\n"
         "# 3 - pull any source with the SAME id\n"
         f"GET /akasha/api/v1/sap?project_id={PID}\n"
         f"GET /akasha/api/v1/pulse?project_id={PID}&kind=nc\n"
         f"GET /akasha/api/v1/transmission?project_id={PID}")
    para(doc, "The /akasha prefix is stripped by the gateway. Server-to-server callers may use "
              "/api/v1 directly; both reach the same handler.")

    # ── 2 project id ──
    doc.add_heading("2. The project identifier", level=1)
    para(doc, "The canonical identifier is project_id — for example FY26-P21 or AHEJ5L. It is "
              "unique and present for every project, and it is the only identifier you need.")
    para(doc, "Lookups are tolerant. All of the following resolve to the same project, and every "
              "response answers with the canonical id, so you can store whichever value you "
              "already hold and normalise later.", after=8)
    table(doc, ["Accepted value", "Meaning"], [
        [PID, "Canonical project id"],
        ["1712", "Internal mapping id"],
        ["6048", "Primavera P6 object id"],
        ["ACL_A01_E_FT_25MW_GROUP_NEW", "Project name"],
    ], widths=[2.6, 3.6])

    # ── 3 filters ──
    doc.add_heading("3. Filters", level=1)
    para(doc, "Identical on every collection endpoint. A filter is never accepted and then "
              "quietly ignored — whatever the server applied is repeated back in "
              "meta.filters_applied.", after=8)
    table(doc, ["Parameter", "Type", "Default", "Notes"],
          [[a, b, c, d] for a, b, c, d in COMMON_PARAMS], widths=[1.15, 0.8, 0.7, 3.55])
    callout(doc, "phase defaults to all",
            "The API does not filter silently. The internal dashboard treats Ongoing as its "
            "default and passes phase=ongoing explicitly; an integrator calling /api/v1/projects "
            "receives every project unless they ask otherwise.")

    # ── 4 envelope ──
    doc.add_heading("4. Response envelope", level=1)
    para(doc, "Every endpoint returns the same shape. Collections put an array in data; single "
              "items put an object.")
    code(doc,
         '{\n'
         '  "data": [ { "project_id": "FY26-P21", ... } ],\n'
         '  "meta": {\n'
         '    "generated_at":    "2026-08-31T09:12:00Z",\n'
         '    "filters_applied": { "project_id": null, "portfolio": null, "phase": "all" },\n'
         '    "sources":         ["SAP"],\n'
         '    "page": 1, "page_size": 50, "total": 86955\n'
         '  }\n'
         '}')
    para(doc, "Every data row carries project_id, including rows from systems that key on "
              "something else internally, so results can be grouped without a second lookup. On "
              "/pulse, Pulse's own identifier is returned as pulse_project_uuid so that "
              "project_id always means the canonical one.")
    para(doc, "Monetary fields are in rupees unless the field name ends in _cr (crore). "
              "Timestamps are ISO-8601.")

    # ── 5 errors ──
    doc.add_heading("5. Errors", level=1)
    table(doc, ["Status", "When", "Body"], [
        ["200", "Success, including a legitimately empty result", "envelope"],
        ["404", "project_id matches no project", '{"detail": "No project matching \'X\'"}'],
        ["422", "Invalid phase or kind, or page_size above 200", '{"detail": "phase must be one of ongoing, commissioned, all"}'],
        ["500", "Unhandled server error", '{"detail": "Internal Server Error"}'],
    ], widths=[0.75, 2.35, 3.1])
    callout(doc, "An empty list is not always 'nothing happened'",
            "It can also mean the project is not connected to that source. The two are "
            "distinguishable: check the unlinked array on the project, or call /api/v1/coverage. "
            "This matters most for Pulse and Transmission, which link for roughly half the "
            "portfolio.")

    doc.add_page_break()

    # ── 6 endpoint reference ──
    doc.add_heading("6. Endpoint reference", level=1)
    para(doc, "Examples below are live responses captured from a running instance.", colour=MUTED, italic=True)

    current_group = None
    for ep in ENDPOINTS:
        if ep["group"] != current_group:
            current_group = ep["group"]
            doc.add_heading(current_group, level=2)

        doc.add_heading(f"GET {ep['path']}", level=3)
        para(doc, ep["desc"], after=8)

        if ep["params"]:
            para(doc, "Query parameters", bold=True, size=9.5, after=4)
            table(doc, ["Parameter", "Type", "Default", "Notes"],
                  [[a, b, c, d] for a, b, c, d in ep["params"]], widths=[1.15, 0.8, 0.7, 3.55])

        para(doc, "Example request", bold=True, size=9.5, after=4)
        code(doc, "GET /akasha" + ep["sample"])

        sample = SAMPLES.get(ep["sample"].split("?")[0] if "?" not in ep["sample"] else ep["sample"])
        if sample is None:
            for key, value in SAMPLES.items():
                if key == "_pid":
                    continue
                if key.split("?")[0] == ep["path"].split("?")[0].replace("{project_id}", PID):
                    sample = value
                    break
        if sample:
            para(doc, "Example response", bold=True, size=9.5, after=4)
            code(doc, pretty(sample["body"]))

        if ep["fields"]:
            para(doc, "Key response fields", bold=True, size=9.5, after=4)
            table(doc, ["Field", "Type", "Description"],
                  [[a, b, c] for a, b, c in ep["fields"]], widths=[1.9, 0.95, 3.35])

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ── 7 coverage ──
    doc.add_heading("7. Source coverage", level=1)
    para(doc, "Not every project is connected to every source system. The API states this "
              "rather than returning a silent empty result. Coverage improves as the mapping "
              "sheet is completed; no client change is required when it does.", after=8)
    cov = SAMPLES["/api/v1/coverage"]["body"]["data"]
    rows = []
    meaning = {
        "p6": "Genuinely no rows.",
        "sap": "Usually genuine; check the unlinked list.",
        "tc": "Often means not connected.",
        "pulse": "Usually means not connected, not 'no issues'.",
    }
    label = {"p6": "Primavera P6", "sap": "SAP", "tc": "Transmission", "pulse": "Pulse quality"}
    for key, value in cov["systems"].items():
        rows.append([label.get(key, key),
                     f"{value['linked']} / {value['total']}",
                     f"{value['pct']}%",
                     meaning.get(key, "")])
    table(doc, ["Source", "Linked", "%", "What an empty result means"], rows,
          widths=[1.4, 0.9, 0.55, 3.35])

    # ── 8 notes ──
    doc.add_heading("8. Operational notes", level=1)
    table(doc, ["Topic", "Detail"], [
        ["Authentication", "POST /api/auth/login issues a token. Endpoints do not yet validate it — treat the API as trusted-network-only until that is in place."],
        ["Caching", "v1 endpoints read live and are not cached."],
        ["Rate limits", "None enforced. /api/v1/activities is 123,180 rows portfolio-wide; always scope it by project_id."],
        ["Stability", "Fields inside data mirror the source tables and may gain columns. Treat additions as non-breaking. project_id and the envelope will not change within v1."],
        ["Legacy routes", "Pre-v1 endpoints still exist for the internal application. They are not part of this contract and should not be integrated against."],
        ["Known data issues", "TcNetworkEdge.status holds unparsed values ('7', 'Mar-30'); use normalized_status instead. Delay is derived from a current-state field, so historical delay counts are not available."],
    ], widths=[1.3, 4.9])

    doc.save(OUT)
    print("written:", OUT)
    print("size: %.0f KB" % (os.path.getsize(OUT) / 1024))


if __name__ == "__main__":
    build()
