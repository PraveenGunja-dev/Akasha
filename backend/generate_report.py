import sys
import os
import re
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from database import SessionLocal
from models import ProjectMapping, P6Activity, P6Project, MTTrialRun, P6WBSNode
import docx

def normalize_block(name):
    name = name.replace(" ", "").upper()
    m = re.match(r'(BLOCK-|WTG-?)0+(\d+)', name)
    if m:
        return f"{m.group(1)}{m.group(2)}"
    return name

def generate_report():
    db = SessionLocal()
    
    doc = docx.Document()
    doc.add_heading('Project Tracker Report', 0)
    
    projects = db.query(ProjectMapping).all()
    
    doc.add_paragraph(f"Total Projects Tracked: {len(projects)}")
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'P6 ID'
    hdr_cells[1].text = 'Project Name'
    hdr_cells[2].text = 'WBS Element'
    hdr_cells[3].text = 'COD Blocks/WTG Done'
    hdr_cells[4].text = 'MW Generated'
    hdr_cells[5].text = 'Pending COD'
    hdr_cells[6].text = 'TR Done but COD Pending'
    
    all_trs = db.query(MTTrialRun).all()
    
    for p in projects:
        wbs = p.module_wbs if p.module_wbs else 'N/A'
        p6_id = p.project_id if p.project_id else 'N/A'
        proj_name = p.project_name_from_p6 if p.project_name_from_p6 else 'N/A'
        
        cod_done = 0
        pending_cod = 0
        tr_done_cod_not = 0
        mw_generated = 0.0
        
        p6_proj = db.query(P6Project).filter(P6Project.project_id == p.project_id).first()
        all_blocks = set()
        if p6_proj:
            obj_id = p6_proj.p6_object_id
            
            # Find all blocks from WBS
            wbs_nodes = db.query(P6WBSNode).filter(P6WBSNode.project_object_id == obj_id).all()
            for w in wbs_nodes:
                m = re.search(r'(Block-\d+|WTG\s*\d+)', w.wbs_name or "", re.IGNORECASE)
                if m:
                    all_blocks.add(normalize_block(m.group(1)))
                    
            # Fetch all COD and TR acts
            cod_acts = db.query(P6Activity).filter(P6Activity.project_object_id == obj_id, P6Activity.name.ilike('%COD%')).all()
            tr_acts = db.query(P6Activity).filter(P6Activity.project_object_id == obj_id, P6Activity.name.ilike('%Trial%')).all()
            
            for a in cod_acts + tr_acts:
                m = re.search(r'(Block-\d+|WTG\s*\d+)', a.name or "", re.IGNORECASE)
                if m:
                    all_blocks.add(normalize_block(m.group(1)))
                    
            blocks_status = {b: {'cod': 'Not Started', 'tr': 'Not Started'} for b in all_blocks}
            
            for a in cod_acts:
                m = re.search(r'(Block-\d+|WTG\s*\d+)', a.name or "", re.IGNORECASE)
                if m:
                    b_name = normalize_block(m.group(1))
                    if a.status == 'Completed':
                        blocks_status[b_name]['cod'] = 'Completed'
            
            for a in tr_acts:
                m = re.search(r'(Block-\d+|WTG\s*\d+)', a.name or "", re.IGNORECASE)
                if m:
                    b_name = normalize_block(m.group(1))
                    if a.status == 'Completed':
                        blocks_status[b_name]['tr'] = 'Completed'
            
            for b, status in blocks_status.items():
                is_cod = (status['cod'] == 'Completed')
                is_tr = (status['tr'] == 'Completed')
                
                if is_cod:
                    cod_done += 1
                else:
                    pending_cod += 1
                    if is_tr:
                        tr_done_cod_not += 1
        
        # Calculate MW generated
        is_solar = any('BLOCK' in b for b in all_blocks)
        is_wind = any('WTG' in b for b in all_blocks)
        
        if is_solar:
            total_blocks = len(all_blocks)
            cap = p.capacity_mwac or p.capacity_mwdc or 0.0
            mw_per_unit = cap / total_blocks if total_blocks > 0 else 0.0
            mw_generated = mw_per_unit * cod_done
        elif is_wind:
            proj_name_nospace = proj_name.replace(" ", "").lower()
            mw_per_wtg = 0.0
            for tr in all_trs:
                tr_proj_name = (tr.project_name_p6 or tr.project_name or "").replace(" ", "").lower()
                if tr_proj_name == proj_name_nospace:
                    if tr.tr_quantity_mw:
                        mw_per_wtg = tr.tr_quantity_mw
                        break
            mw_generated = mw_per_wtg * cod_done
        else:
            proj_name_nospace = proj_name.replace(" ", "").lower()
            for tr in all_trs:
                tr_proj_name = (tr.project_name_p6 or tr.project_name or "").replace(" ", "").lower()
                if tr_proj_name == proj_name_nospace:
                    if tr.activity_name and 'COD' in tr.activity_name.upper():
                        mw_generated += (tr.tr_quantity_mw or 0.0)
                        
        row_cells = table.add_row().cells
        row_cells[0].text = str(p6_id)
        row_cells[1].text = str(proj_name)
        row_cells[2].text = str(wbs)
        row_cells[3].text = str(cod_done)
        row_cells[4].text = str(round(mw_generated, 2))
        row_cells[5].text = str(pending_cod)
        row_cells[6].text = str(tr_done_cod_not)
        
    db.close()
    
    output_path = os.path.join(os.path.dirname(__file__), "project_tracker_report_v7.docx")
    doc.save(output_path)
    print(f"Report successfully generated at: {output_path}")

if __name__ == "__main__":
    generate_report()
