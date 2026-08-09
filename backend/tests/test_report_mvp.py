import os
import sys
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
import unittest
from unittest.mock import patch

from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker


os.environ["DATABASE_URL"] = "sqlite:///:memory:"
BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from database import Base
import models
from services.report_mvp_service import (
    build_project_comparison_dataset,
    build_project_progress_dataset,
    build_portfolio_progress_dataset,
    create_preview_token,
    create_project_comparison_preview,
    generate_narrative,
    generate_project_comparison_report,
    validate_preview_token,
)
from services.report_renderers import (
    _capacity_rows,
    _risk_rows,
    render_project_progress_docx,
    render_project_progress_pdf,
    render_portfolio_progress_docx,
    render_portfolio_progress_pdf,
    render_visualization_spec,
)
from services.report_visualization_service import (
    resolve_visualization_references,
    select_conversation_visualizations,
)


def report_dataset():
    return {
        "metadata": {
            "project_name": "Test Project",
            "source_freshness": {"P6": "2026-07-22", "SAP": None, "TC": None, "Pulse": None},
            "missing_sources": ["SAP", "TC", "Pulse"],
        },
        "executive_summary": "Test Project is 23.1% complete. SPI is unavailable.",
        "project_summary": {
            "status": "Active",
            "scheduled_finish": "2027-10-01",
            "data_date": "2026-07-18",
            "last_synced_at": "2026-07-22",
        },
        "schedule": {
            "progress_pct": 23.1,
            "completed_activities": 486,
            "in_progress_activities": 84,
            "not_started_activities": 1730,
            "spi": None,
            "cpi": None,
        },
        "procurement": {"has_data": False, "reason": "No mapped data"},
        "transmission": {"has_data": False, "reason": "No mapped data"},
        "quality": {"has_data": False, "non_conformances": 0, "rfis": 0},
        "capacity": {
            "projects": [{
                "total_capacity": 100,
                "cod_mw": 25,
                "tr_mw": 12.5,
                "remaining_capacity": 62.5,
            }],
            "metadata": {"formula": {"version": "dashboard-capacity-overview-v1"}},
        },
        "risk": {
            "project360.status_tier": {
                "metric_id": "project360.status_tier",
                "name": "Project360 status tier",
                "value": "Watchlist",
            },
        },
        "report_visualizations": {
            "daily_completion_trend": {
                "daily": [
                    {"date": "2026-07-30", "activities_completed": 1, "cumulative_activity_finish_pct": 20},
                    {"date": "2026-07-31", "activities_completed": 2, "cumulative_activity_finish_pct": 30},
                ],
            },
            "block_progress": {
                "blocks": [
                    {"block": "BLOCK-01", "current_activity_completion_pct": 75},
                    {"block": "BLOCK-02", "current_activity_completion_pct": 42},
                ],
            },
        },
        "in_progress_activities": {
            "activities": [{"activity_id": "A-1", "name": "Foundation", "percent_complete": 75.0}],
        },
    }


class ReportMvpTests(unittest.TestCase):
    def test_conversation_chart_selection_respects_scope_topic_and_explicit_choice(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        db = sessionmaker(bind=engine)()
        db.add(models.ChatSession(session_id="chart-session", owner_subject="user", tenant_id="tenant"))
        db.add_all([
            models.ChatMessage(
                session_id="chart-session", role="assistant", content="schedule", status="completed",
                project_ids="P-1", visualizations=[{
                    "title": "Schedule delay by block", "report_inclusion": "auto",
                    "spec": {"schema_version": "visualization.v2", "shape": "bar", "title": "Schedule delay by block"},
                }],
            ),
            models.ChatMessage(
                session_id="chart-session", role="assistant", content="quality", status="completed",
                project_ids="P-1", visualizations=[{
                    "title": "Quality issues", "report_inclusion": "auto",
                    "spec": {"schema_version": "visualization.v2", "shape": "bar", "title": "Quality issues"},
                }],
            ),
            models.ChatMessage(
                session_id="chart-session", role="assistant", content="other", status="completed",
                project_ids="P-2", visualizations=[{
                    "title": "Other project capacity", "report_inclusion": "include",
                    "spec": {"schema_version": "visualization.v2", "shape": "bar", "title": "Other project capacity"},
                }],
            ),
        ])
        db.commit()

        selected, excluded = select_conversation_visualizations(
            db, session_id="chart-session", scope_kind="project", scope_project_ids=["P-1"],
            selection_text="Include only schedule charts in the report",
        )
        self.assertEqual(
            [item.report_payload()["title"] for item in selected],
            ["Schedule delay by block", "Other project capacity"],
        )
        self.assertTrue(any(item["title"] == "Quality issues" for item in excluded))
        without_quality, _ = select_conversation_visualizations(
            db, session_id="chart-session", scope_kind="project", scope_project_ids=["P-1"],
            selection_text="Create the report without quality charts",
        )
        self.assertEqual(
            [item.report_payload()["title"] for item in without_quality],
            ["Schedule delay by block", "Other project capacity"],
        )
        refs = [item.reference() for item in selected]
        resolved = resolve_visualization_references(db, session_id="chart-session", references=refs)
        self.assertEqual([item["snapshot_hash"] for item in resolved], [item["h"] for item in refs])
        db.close()
        engine.dispose()

    def test_v2_chart_renders_and_invalid_chart_falls_back_in_both_reports(self):
        dataset = report_dataset()
        v2 = {
            "schema_version": "visualization.v2", "chart_id": "chat-progress", "chart_type": "bar",
            "shape": "bar", "title": "Saved Progress", "subtitle": "Exact chat snapshot",
            "summary": "Progress saved in chat.", "accessibility_description": "Progress bars.",
            "encoding": {
                "x": {"field": "block", "label": "Block", "field_type": "categorical"},
                "y": [{"field": "progress", "label": "Progress", "field_type": "quantitative", "value_format": "percent"}],
            },
            "data": [{"block": "A", "progress": 55}, {"block": "B", "progress": 72}],
            "source_tables": ["p6_activity"], "spec_hash": "sha256:" + "1" * 64,
        }
        self.assertGreater(len(render_visualization_spec(v2).getvalue()), 1_000)
        dataset["conversation_visualizations"] = [
            {"title": "Saved Progress", "summary": "Progress saved in chat.", "report_section": "schedule", "spec": v2},
            {
                "title": "Unsupported saved chart", "report_section": "appendix",
                "spec": {"schema_version": "visualization.v2", "shape": "future_shape"},
                "data_table": [{"item": "A", "value": 10}],
            },
        ]
        with TemporaryDirectory() as directory:
            pdf = Path(directory) / "conversation.pdf"
            docx = Path(directory) / "conversation.docx"
            render_project_progress_pdf(dataset, pdf)
            render_project_progress_docx(dataset, docx)
            self.assertGreater(pdf.stat().st_size, 2_000)
            from docx import Document
            text = "\n".join(paragraph.text for paragraph in Document(docx).paragraphs)
            self.assertIn("Saved Progress", text)
            self.assertIn("saved data snapshot is shown instead", text)

    def test_dataset_supports_catalog_project_without_p6(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        db = sessionmaker(bind=engine)()
        db.add(models.ProjectMapping(
            project_id="MAP-1",
            project="Mapping Only",
            project_name_from_p6="Mapping Only Project",
            capacity_mwac=100,
        ))
        db.commit()

        dataset = build_project_progress_dataset(db, "MAP-1")

        self.assertEqual(dataset["metadata"]["project_name"], "Mapping Only Project")
        self.assertFalse(dataset["project_summary"]["p6_available"])
        self.assertIsNone(dataset["schedule"]["progress_pct"])
        self.assertIn("P6", dataset["metadata"]["missing_sources"])
        self.assertEqual(dataset["in_progress_activities"]["activities"], [])
        with patch.dict(os.environ, {"AKASHA_REPORT_AI_NARRATIVE": "false"}):
            self.assertIn("P6 schedule facts are unavailable", generate_narrative(dataset))
        db.close()
        engine.dispose()

    def test_preview_token_is_bound_to_session_and_project(self):
        token = create_preview_token("session-a", "FY26-P18")
        validate_preview_token(token, "session-a", "FY26-P18")
        with self.assertRaises(ValueError):
            validate_preview_token(token, "session-b", "FY26-P18")
        with self.assertRaises(ValueError):
            validate_preview_token(token, "session-a", "FY26-P19")

    def test_pdf_and_docx_render_from_same_dataset(self):
        with TemporaryDirectory() as directory:
            pdf = Path(directory) / "report.pdf"
            docx = Path(directory) / "report.docx"
            dataset = report_dataset()
            render_project_progress_pdf(dataset, pdf)
            render_project_progress_docx(dataset, docx)
            self.assertGreater(pdf.stat().st_size, 1_000)
            self.assertGreater(docx.stat().st_size, 10_000)
            self.assertEqual(pdf.read_bytes()[:4], b"%PDF")
            self.assertEqual(docx.read_bytes()[:2], b"PK")
            from docx import Document
            text = "\n".join(paragraph.text for paragraph in Document(docx).paragraphs)
            self.assertIn("executive reading layer", text)
            self.assertIn("Executive takeaway:", text)
            self.assertIn("BLOCK-01 has the highest displayed completion", text)

    def test_capacity_and_named_risk_sections_use_authoritative_facts(self):
        dataset = report_dataset()
        self.assertIn(("COD", "25 MW"), _capacity_rows(dataset["capacity"]))
        self.assertIn(
            ("Project360 status tier", "Watchlist"),
            _risk_rows(dataset["risk"]),
        )

    def test_portfolio_report_dataset_and_renderers(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        db = sessionmaker(bind=engine)()
        db.add(models.ProjectMapping(
            project_id="PORT-1",
            project="Portfolio Project",
            project_name_from_p6="Portfolio Project P6",
        ))
        db.add(models.P6Project(
            p6_object_id=901,
            project_id="PORT-1",
            duration_percent_complete=0.6,
        ))
        db.commit()

        dataset = build_portfolio_progress_dataset(db)
        dataset["executive_summary"] = "A grounded portfolio summary suitable for rendering."
        self.assertEqual(dataset["summary"]["total_projects"], 1)
        self.assertEqual(dataset["metadata"]["period"], "current_month")
        self.assertEqual(dataset["report_visualizations"]["project_progress"]["schema_version"], "visualization.v1")
        self.assertEqual(dataset["report_visualizations"]["schedule_status"]["shape"], "donut")
        with TemporaryDirectory() as directory:
            pdf = Path(directory) / "portfolio.pdf"
            docx = Path(directory) / "portfolio.docx"
            render_portfolio_progress_pdf(dataset, pdf)
            render_portfolio_progress_docx(dataset, docx)
            self.assertEqual(pdf.read_bytes()[:4], b"%PDF")
            self.assertEqual(docx.read_bytes()[:2], b"PK")
        db.close()
        engine.dispose()

    def test_comparison_report_supports_preview_and_direct_generation(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        db = sessionmaker(bind=engine)()
        db.add_all([
            models.ProjectMapping(
                project_id="CMP-1", project="Comparison One",
                project_name_from_p6="Comparison One", capacity_mwac=125, spv_name="SPV One",
            ),
            models.ProjectMapping(
                project_id="CMP-2", project="Comparison Two",
                project_name_from_p6="Comparison Two", capacity_mwac=150, spv_name="SPV Two",
            ),
            models.P6Project(
                p6_object_id=1001, project_id="CMP-1", duration_percent_complete=0.91,
                planned_duration=4000, actual_duration=3500, remaining_duration=500,
                baseline_finish_date=__import__("datetime").datetime(2026, 6, 1),
                finish_date=__import__("datetime").datetime(2026, 8, 1),
                data_date=__import__("datetime").datetime(2026, 7, 1),
            ),
            models.P6Project(
                p6_object_id=1002, project_id="CMP-2", duration_percent_complete=0.94,
                planned_duration=4800, actual_duration=3900, remaining_duration=900,
                baseline_finish_date=__import__("datetime").datetime(2026, 7, 1),
                finish_date=__import__("datetime").datetime(2026, 9, 15),
                data_date=__import__("datetime").datetime(2026, 7, 8),
            ),
        ])
        db.commit()
        runtime = SimpleNamespace(
            session_id="comparison-session", user_id="user-1", tenant_id="tenant-1"
        )
        preview = create_project_comparison_preview(db, runtime, ["CMP-1", "CMP-2"])
        self.assertEqual(preview["status"], "awaiting_confirmation")
        self.assertEqual(preview["formats"], ["PDF", "DOCX"])
        dataset = build_project_comparison_dataset(db, ["CMP-1", "CMP-2"])
        self.assertEqual(
            [spec["shape"] for spec in dataset["report_visualizations"].values()],
            ["radial_progress", "horizontal_bar", "vertical_bar", "lollipop"],
        )
        with TemporaryDirectory() as directory, patch.dict(
            os.environ, {"AKASHA_REPORT_ARTIFACT_DIR": directory}
        ):
            generated = generate_project_comparison_report(
                db, runtime, ["CMP-1", "CMP-2"], preview["preview_token"]
            )
            self.assertEqual(generated["status"], "generated")
            self.assertEqual([row["format"] for row in generated["downloads"]], ["PDF", "DOCX"])
            self.assertTrue(all((Path(directory) / row["filename"]).exists() for row in generated["downloads"]))
            direct = generate_project_comparison_report(db, runtime, ["CMP-1", "CMP-2"])
            self.assertEqual(direct["status"], "generated")
            self.assertEqual([row["format"] for row in direct["downloads"]], ["PDF", "DOCX"])
        db.close()
        engine.dispose()

    def test_valid_structured_narrative_is_used(self):
        expected = (
            "Test Project remains active with P6 duration progress at 23.1%. "
            "SPI and CPI are unavailable, so no schedule-performance classification is made. "
            "SAP, TC, and Pulse data are not mapped and remain unassessed."
        )
        provider = SimpleNamespace(invoke=lambda *_args, **_kwargs: SimpleNamespace(
            content='{"executive_summary": ' + __import__("json").dumps(expected) + '}'
        ))
        with patch("services.report_mvp_service.get_model_provider", return_value=provider):
            self.assertEqual(generate_narrative(report_dataset()), expected)

    def test_reasoning_leak_falls_back_to_clean_deterministic_summary(self):
        leaked = (
            "The user wants a concise paragraph. I need to analyze the supplied JSON facts. "
            "Let me draft the report after reviewing the constraints and available metrics."
        )
        provider = SimpleNamespace(invoke=lambda *_args, **_kwargs: SimpleNamespace(
            content='{"executive_summary": ' + __import__("json").dumps(leaked) + '}'
        ))
        with patch("services.report_mvp_service.get_model_provider", return_value=provider):
            result = generate_narrative(report_dataset())
        self.assertNotIn("user wants", result.lower())
        self.assertNotIn("i need to", result.lower())
        self.assertIn("23.1%", result)
        self.assertIn("No mapped data is available from SAP, TC, Pulse", result)


if __name__ == "__main__":
    unittest.main()
