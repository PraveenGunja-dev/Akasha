from __future__ import annotations

from io import BytesIO
from pathlib import Path


REPORT_BLUE = "#1769AA"
REPORT_TEAL = "#0891B2"
REPORT_GREEN = "#059669"
REPORT_AMBER = "#D97706"
REPORT_RED = "#DC2626"


def _font(size: int, *, bold: bool = False):
    from PIL import ImageFont

    candidates = (
        ["C:/Windows/Fonts/arialbd.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"]
        if bold else
        ["C:/Windows/Fonts/arial.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"]
    )
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def _chart_canvas(title: str, subtitle: str, width: int = 1500, height: int = 760):
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((20, 20, width - 20, height - 20), radius=24, fill="#FFFFFF", outline="#D7E1EA", width=2)
    draw.text((64, 48), title, fill="#172033", font=_font(30, bold=True))
    draw.text((64, 90), subtitle, fill="#667085", font=_font(18))
    return image, draw


def _png_bytes(image) -> BytesIO:
    stream = BytesIO()
    image.save(stream, format="PNG", optimize=True)
    stream.seek(0)
    return stream


def _series_color(series: dict, index: int | None = None) -> str:
    from services.visualization_spec import semantic_color

    if index is not None:
        item_colors = series.get("item_semantic_colors") or []
        if index < len(item_colors):
            return semantic_color(item_colors[index])
    return semantic_color(str(series.get("semantic_color") or "primary"))


def render_visualization_spec(spec: dict | None) -> BytesIO | None:
    """Render a validated declarative visualization without executing chart code."""
    if not spec:
        return None
    if spec.get("schema_version") == "visualization.v2":
        return _render_visualization_spec_v2(spec)
    shape = spec.get("shape")
    categories = list(spec.get("categories") or [])
    series = list(spec.get("series") or [])
    if not categories or not series:
        return None
    title = str(spec.get("title") or "Akasha visualization")
    subtitle = str(spec.get("subtitle") or "")

    if shape == "horizontal_bar":
        values = list(series[0].get("values") or [])
        rows = [(label, float(value or 0)) for label, value in zip(categories, values)][:12]
        image, draw = _chart_canvas(title, subtitle, height=max(560, 180 + len(rows) * 46))
        left, right, top = 420, 1400, 150
        chart_width = right - left
        maximum = max((value for _, value in rows), default=1) or 1
        value_format = series[0].get("value_format")
        for index, (label, value) in enumerate(rows):
            y = top + index * 46
            draw.text((64, y + 4), label[:38], fill="#344054", font=_font(17))
            draw.rounded_rectangle((left, y, right, y + 26), radius=8, fill="#EDF2F7")
            filled = max(4, int(chart_width * max(0, value) / maximum))
            draw.rounded_rectangle(
                (left, y, left + filled, y + 26), radius=8,
                fill=_series_color(series[0], index),
            )
            suffix = "%" if value_format == "percent" else "d" if value_format == "days" else ""
            decimals = 1 if value_format in {"percent", "decimal"} else 0
            label_value = f"{value:.{decimals}f}{suffix}"
            draw.text((min(left + filled + 12, right - 76), y + 2), label_value, fill="#172033", font=_font(17, bold=True))
        return _png_bytes(image)

    if shape == "donut":
        values = [int(value or 0) for value in (series[0].get("values") or [])]
        total = sum(values)
        if not total:
            return None
        image, draw = _chart_canvas(title, subtitle, width=1100, height=680)
        bounds = (100, 160, 570, 630)
        start = -90
        for index, value in enumerate(values):
            if value:
                sweep = 360 * value / total
                draw.pieslice(bounds, start=start, end=start + sweep, fill=_series_color(series[0], index))
                start += sweep
        draw.ellipse((205, 265, 465, 525), fill="#FFFFFF")
        draw.text((335, 340), str(total), fill="#172033", font=_font(46, bold=True), anchor="mm")
        draw.text((335, 390), str(series[0].get("name") or "items").lower(), fill="#667085", font=_font(17), anchor="mm")
        y = 190
        for index, (label, value) in enumerate(zip(categories, values)):
            draw.rounded_rectangle((650, y, 672, y + 22), radius=5, fill=_series_color(series[0], index))
            draw.text((692, y - 1), f"{label}: {value}", fill="#344054", font=_font(19, bold=True))
            y += 62
        return _png_bytes(image)

    if shape == "combo" and len(series) >= 2:
        bars = list(series[0].get("values") or [])
        line = list(series[1].get("values") or [])
        image, draw = _chart_canvas(title, subtitle)
        left, right, top, bottom = 90, 1410, 150, 650
        draw.line((left, bottom, right, bottom), fill="#98A2B3", width=2)
        maximum = max((float(value or 0) for value in bars), default=1) or 1
        points = []
        step = (right - left) / max(1, len(categories))
        for index, value in enumerate(bars):
            x = left + index * step + step * 0.18
            bar_height = (bottom - top) * float(value or 0) / maximum
            draw.rounded_rectangle(
                (x, bottom - bar_height, x + step * 0.55, bottom), radius=4,
                fill=_series_color(series[0]),
            )
            if index < len(line) and line[index] is not None:
                points.append((x + step * 0.28, bottom - (bottom - top) * float(line[index]) / 100))
        if len(points) > 1:
            line_color = _series_color(series[1])
            draw.line(points, fill=line_color, width=5, joint="curve")
            for x, y in points[::max(1, len(points) // 10)]:
                draw.ellipse((x - 5, y - 5, x + 5, y + 5), fill=line_color)
        draw.text((left, 684), categories[0], fill="#667085", font=_font(16))
        draw.text((right - 120, 684), categories[-1], fill="#667085", font=_font(16))
        return _png_bytes(image)

    if shape == "vertical_bar":
        image, draw = _chart_canvas(title, subtitle)
        left, right, top, bottom = 110, 1410, 150, 640
        draw.line((left, bottom, right, bottom), fill="#98A2B3", width=2)
        all_values = [float(value or 0) for item in series for value in (item.get("values") or [])]
        maximum = max(all_values, default=1) or 1
        group_width = (right - left) / max(1, len(categories))
        bar_width = min(70, group_width * 0.68 / max(1, len(series)))
        for category_index, category in enumerate(categories):
            group_left = left + category_index * group_width + group_width * 0.16
            for series_index, item in enumerate(series):
                values = item.get("values") or []
                value = float(values[category_index] or 0) if category_index < len(values) else 0
                height = (bottom - top) * value / maximum
                x = group_left + series_index * bar_width
                draw.rounded_rectangle(
                    (x, bottom - height, x + bar_width - 5, bottom), radius=5,
                    fill=_series_color(item),
                )
            draw.text(
                (left + category_index * group_width + group_width / 2, bottom + 20),
                category[:22], fill="#344054", font=_font(14), anchor="ma",
            )
        legend_x = left
        for item in series:
            draw.rounded_rectangle((legend_x, 690, legend_x + 22, 712), radius=4, fill=_series_color(item))
            draw.text((legend_x + 30, 690), str(item.get("name")), fill="#344054", font=_font(15))
            legend_x += 210
        return _png_bytes(image)

    if shape == "radial_progress":
        image, draw = _chart_canvas(title, subtitle)
        values = list(series[0].get("values") or [])
        count = max(1, len(categories))
        for index, (category, raw_value) in enumerate(zip(categories, values)):
            value = max(0.0, min(100.0, float(raw_value or 0)))
            center_x = int((index + 0.5) * 1500 / count)
            bounds = (center_x - 150, 210, center_x + 150, 510)
            draw.arc(bounds, start=140, end=400, fill="#E4EAF0", width=30)
            draw.arc(bounds, start=140, end=140 + 260 * value / 100, fill=_series_color(series[0]), width=30)
            draw.text((center_x, 350), f"{value:.1f}%", fill="#172033", font=_font(38, bold=True), anchor="mm")
            draw.text((center_x, 560), category[:28], fill="#344054", font=_font(17, bold=True), anchor="mm")
        return _png_bytes(image)

    if shape == "lollipop":
        values = [float(value or 0) for value in (series[0].get("values") or [])]
        image, draw = _chart_canvas(title, subtitle, height=max(560, 220 + len(categories) * 90))
        left, right, top = 420, 1400, 190
        minimum = min([0.0, *values])
        maximum = max([1.0, *values])
        span = maximum - minimum or 1
        zero_x = left + (0 - minimum) / span * (right - left)
        draw.line((zero_x, top - 30, zero_x, top + len(categories) * 90), fill="#98A2B3", width=2)
        for index, (category, value) in enumerate(zip(categories, values)):
            y = top + index * 90
            x = left + (value - minimum) / span * (right - left)
            item_color = _series_color(series[0], index)
            draw.text((64, y - 10), category[:38], fill="#344054", font=_font(18))
            draw.line((zero_x, y, x, y), fill=item_color, width=6)
            draw.ellipse((x - 12, y - 12, x + 12, y + 12), fill=item_color)
            draw.text((x + 22, y - 12), f"{value:.0f}d", fill="#172033", font=_font(18, bold=True))
        return _png_bytes(image)
    return None


def _render_visualization_spec_v2(spec: dict) -> BytesIO | None:
    rows = list(spec.get("data") or [])[:100]
    encoding = spec.get("encoding") or {}
    x_channel = encoding.get("x") or encoding.get("label") or {}
    y_channels = list(encoding.get("y") or [])
    if not rows or not y_channels:
        return None
    x_field = x_channel.get("field")
    if not x_field:
        return None
    title = str(spec.get("title") or "Conversation visualization")
    subtitle = str(spec.get("subtitle") or spec.get("summary") or "")[:150]
    shape = spec.get("shape")
    palette = [REPORT_BLUE, REPORT_TEAL, REPORT_GREEN, REPORT_AMBER, REPORT_RED, "#7C3AED"]

    def numeric(value):
        try:
            return float(value)
        except (TypeError, ValueError):
            return 0.0

    if shape in {"bar", "horizontal_bar", "stacked_bar", "waterfall", "donut"}:
        categories = [str(row.get(x_field, "")) for row in rows][:20]
        series = [{
            "name": str(channel.get("label") or channel.get("field")),
            "shape": "donut" if shape == "donut" else "bar",
            "values": [numeric(row.get(channel.get("field"))) for row in rows[:20]],
            "semantic_color": "primary" if index == 0 else "success" if index == 1 else "warning",
            "value_format": channel.get("value_format") or "decimal",
            "axis_index": int(channel.get("axis_index") or 0),
        } for index, channel in enumerate(y_channels[:4])]
        compatible = {
            **spec,
            "schema_version": "visualization.v1",
            "shape": "donut" if shape == "donut" else "horizontal_bar" if shape == "horizontal_bar" else "vertical_bar",
            "categories": categories,
            "series": series,
        }
        return render_visualization_spec(compatible)

    if shape in {"line", "scatter"}:
        image, draw = _chart_canvas(title, subtitle)
        left, right, top, bottom = 110, 1410, 160, 640
        draw.line((left, bottom, right, bottom), fill="#98A2B3", width=2)
        values = [numeric(row.get(channel.get("field"))) for channel in y_channels[:4] for row in rows]
        minimum, maximum = min([0.0, *values]), max([1.0, *values])
        span = maximum - minimum or 1
        step = (right - left) / max(1, len(rows) - 1)
        for series_index, channel in enumerate(y_channels[:4]):
            points = [
                (left + row_index * step, bottom - (numeric(row.get(channel.get("field"))) - minimum) / span * (bottom - top))
                for row_index, row in enumerate(rows)
            ]
            color = palette[series_index]
            if shape == "line" and len(points) > 1:
                draw.line(points, fill=color, width=5, joint="curve")
            for x, y in points:
                draw.ellipse((x - 6, y - 6, x + 6, y + 6), fill=color)
            legend_x = left + series_index * 260
            draw.rectangle((legend_x, 690, legend_x + 20, 710), fill=color)
            draw.text((legend_x + 28, 688), str(channel.get("label") or channel.get("field"))[:24], fill="#344054", font=_font(15))
        draw.text((left, 654), str(rows[0].get(x_field, ""))[:24], fill="#667085", font=_font(14))
        draw.text((right, 654), str(rows[-1].get(x_field, ""))[:24], fill="#667085", font=_font(14), anchor="ra")
        return _png_bytes(image)

    if shape == "heatmap":
        color_channel = encoding.get("color") or {}
        color_field = color_channel.get("field") or y_channels[0].get("field")
        y_field = y_channels[0].get("field")
        x_values = list(dict.fromkeys(str(row.get(x_field, "")) for row in rows))[:12]
        y_values = list(dict.fromkeys(str(row.get(y_field, "")) for row in rows))[:10]
        if not x_values or not y_values:
            return None
        heat_values = [numeric(row.get(color_field)) for row in rows]
        maximum = max(heat_values, default=1) or 1
        image, draw = _chart_canvas(title, subtitle, height=max(620, 230 + len(y_values) * 48))
        left, top, cell_w, cell_h = 300, 170, min(85, 1050 / len(x_values)), 45
        lookup = {(str(row.get(x_field, "")), str(row.get(y_field, ""))): numeric(row.get(color_field)) for row in rows}
        for yi, y_label in enumerate(y_values):
            draw.text((64, top + yi * cell_h + 12), y_label[:28], fill="#344054", font=_font(15))
            for xi, x_label in enumerate(x_values):
                value = lookup.get((x_label, y_label), 0)
                intensity = int(235 - 165 * max(0, value) / maximum)
                color = (intensity, min(220, intensity + 35), 245)
                bounds = (left + xi * cell_w, top + yi * cell_h, left + (xi + 1) * cell_w - 3, top + (yi + 1) * cell_h - 3)
                draw.rectangle(bounds, fill=color)
                draw.text(((bounds[0] + bounds[2]) / 2, bounds[1] + 10), f"{value:g}", fill="#172033", font=_font(13), anchor="ma")
        for xi, label in enumerate(x_values):
            draw.text((left + xi * cell_w + cell_w / 2, top + len(y_values) * cell_h + 8), label[:10], fill="#667085", font=_font(12), anchor="ma")
        return _png_bytes(image)
    return None


def _horizontal_bar_chart(rows: list[tuple[str, float]], title: str, subtitle: str) -> BytesIO | None:
    if not rows:
        return None
    rows = rows[:12]
    image, draw = _chart_canvas(title, subtitle, height=max(560, 180 + len(rows) * 46))
    left, right, top = 420, 1400, 150
    chart_width = right - left
    maximum = max((value for _, value in rows), default=1) or 1
    for index, (label, value) in enumerate(rows):
        y = top + index * 46
        draw.text((64, y + 4), label[:38], fill="#344054", font=_font(17))
        draw.rounded_rectangle((left, y, right, y + 26), radius=8, fill="#EDF2F7")
        filled = max(4, int(chart_width * max(0, value) / maximum))
        color = REPORT_GREEN if value >= 75 else REPORT_BLUE if value >= 40 else REPORT_AMBER
        draw.rounded_rectangle((left, y, left + filled, y + 26), radius=8, fill=color)
        draw.text((min(left + filled + 12, right - 62), y + 2), f"{value:.1f}%", fill="#172033", font=_font(17, bold=True))
    return _png_bytes(image)


def _status_donut_chart(counts: dict, title: str, subtitle: str) -> BytesIO | None:
    values = [
        ("Delayed", int(counts.get("delayed", 0)), REPORT_RED),
        ("On track", int(counts.get("on_track", 0)), REPORT_GREEN),
        ("Completed", int(counts.get("completed", 0)), REPORT_BLUE),
        ("P6 unavailable", int(counts.get("p6_unavailable", 0)), "#98A2B3"),
    ]
    total = sum(value for _, value, _ in values)
    if not total:
        return None
    image, draw = _chart_canvas(title, subtitle, width=1100, height=680)
    bounds = (100, 160, 570, 630)
    start = -90
    for label, value, color in values:
        if not value:
            continue
        sweep = 360 * value / total
        draw.pieslice(bounds, start=start, end=start + sweep, fill=color)
        start += sweep
    draw.ellipse((205, 265, 465, 525), fill="#FFFFFF")
    draw.text((302, 340), str(total), fill="#172033", font=_font(46, bold=True), anchor="mm")
    draw.text((335, 390), "projects", fill="#667085", font=_font(17), anchor="mm")
    y = 190
    for label, value, color in values:
        draw.rounded_rectangle((650, y, 672, y + 22), radius=5, fill=color)
        draw.text((692, y - 1), f"{label}: {value}", fill="#344054", font=_font(19, bold=True))
        y += 62
    return _png_bytes(image)


def _daily_trend_chart(data: dict) -> BytesIO | None:
    rows = data.get("daily") or []
    if not rows:
        return None
    image, draw = _chart_canvas(
        "Daily Completion Trend",
        "Activity actual-finish events; not historical duration-percent progress",
    )
    left, right, top, bottom = 90, 1410, 150, 650
    draw.line((left, bottom, right, bottom), fill="#98A2B3", width=2)
    maximum = max((row.get("activities_completed") or 0 for row in rows), default=1) or 1
    points = []
    step = (right - left) / max(1, len(rows))
    for index, row in enumerate(rows):
        x = left + index * step + step * 0.18
        count = row.get("activities_completed") or 0
        bar_height = (bottom - top) * count / maximum
        draw.rounded_rectangle((x, bottom - bar_height, x + step * 0.55, bottom), radius=4, fill=REPORT_TEAL)
        cumulative = row.get("cumulative_activity_finish_pct")
        if cumulative is not None:
            points.append((x + step * 0.28, bottom - (bottom - top) * cumulative / 100))
    if len(points) > 1:
        draw.line(points, fill=REPORT_BLUE, width=5, joint="curve")
        for x, y in points[::max(1, len(points) // 10)]:
            draw.ellipse((x - 5, y - 5, x + 5, y + 5), fill=REPORT_BLUE)
    draw.text((left, 684), str(rows[0].get("date")), fill="#667085", font=_font(16))
    draw.text((right - 100, 684), str(rows[-1].get("date")), fill="#667085", font=_font(16))
    return _png_bytes(image)


def _project_report_charts(dataset: dict) -> list[BytesIO]:
    visualizations = dataset.get("report_visualizations") or {}
    images = []
    trend_spec = visualizations.get("daily_completion_trend") or {}
    trend = (
        render_visualization_spec(trend_spec)
        if trend_spec.get("schema_version") == "visualization.v1"
        else _daily_trend_chart(trend_spec)
    )
    if trend:
        images.append(trend)
    block_spec = visualizations.get("block_progress") or {}
    if block_spec.get("schema_version") == "visualization.v1":
        block = render_visualization_spec(block_spec)
    else:
        block_rows = [
            (row["block"], float(row["current_activity_completion_pct"]))
            for row in (block_spec.get("blocks") or [])
            if row.get("current_activity_completion_pct") is not None
        ]
        block_rows.sort(key=lambda row: row[1], reverse=True)
        block = _horizontal_bar_chart(
            block_rows,
            "Block Progress Snapshot",
            "Current average activity completion by BLOCK-* WBS branch",
        )
    if block:
        images.append(block)
    return images


def _configure_docx(document, *, landscape: bool = False, running_label: str = "Project Progress Report") -> None:
    """Apply the standard_business_brief token set with one landscape report override."""
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    section = document.sections[0]
    section.page_width = Inches(11 if landscape else 8.5)
    section.page_height = Inches(8.5 if landscape else 11)
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    margin = Inches(0.65 if landscape else 1.0)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = running_label
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.runs[0]
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(8.5)
    header_run.font.color.rgb = RGBColor(102, 112, 133)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Confidential | Akasha")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = RGBColor(102, 112, 133)

    # Quiet header rule, matching the memo_masthead family without title decoration.
    p_pr = header._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for key, value in (("val", "single"), ("sz", "4"), ("space", "4"), ("color", "D7E1EA")):
        bottom.set(qn(f"w:{key}"), value)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _style_docx_table(table, widths: list[float] | None = None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    table.autofit = False
    table.style = "Table Grid"
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
            if row_index == 0:
                shading = tc_pr.find(qn("w:shd"))
                if shading is None:
                    shading = OxmlElement("w:shd")
                    tc_pr.append(shading)
                shading.set(qn("w:fill"), "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 77, 120)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0


def _rows(dataset: dict) -> list[tuple[str, str]]:
    summary = dataset["project_summary"]
    schedule = dataset["schedule"]
    return [
        ("Project status", str(summary.get("status") or "Unavailable")),
        ("P6 duration progress", f"{schedule.get('progress_pct')}%" if schedule.get("progress_pct") is not None else "Unavailable"),
        ("Completed activities", str(schedule.get("completed_activities"))),
        ("In-progress activities", str(schedule.get("in_progress_activities"))),
        ("Not-started activities", str(schedule.get("not_started_activities"))),
        ("SPI", str(schedule.get("spi")) if schedule.get("spi") is not None else "Unavailable"),
        ("CPI", str(schedule.get("cpi")) if schedule.get("cpi") is not None else "Unavailable"),
        ("Scheduled finish", str(summary.get("scheduled_finish") or "Unavailable")),
        ("P6 data date", str(summary.get("data_date") or "Unavailable")),
        ("P6 last synchronized", str(summary.get("last_synced_at") or "Unavailable")),
    ]


def _domain_rows(value: dict, limit: int = 10) -> list[tuple[str, str]]:
    rows = []
    for key, item in (value or {}).items():
        if key.startswith("_") or isinstance(item, (dict, list)) or item is None:
            continue
        rows.append((key.replace("_", " ").title(), str(item)))
        if len(rows) >= limit:
            break
    return rows or [("Status", "No mapped data")]


def _capacity_rows(value: dict) -> list[tuple[str, str]]:
    projects = value.get("projects") or []
    if not projects:
        return [("Status", "No mapped data")]
    project = projects[0]
    return [
        ("Total Capacity", f"{project.get('total_capacity')} MW"),
        ("COD", f"{project.get('cod_mw')} MW"),
        ("Trial Run", f"{project.get('tr_mw')} MW"),
        ("Remaining", f"{project.get('remaining_capacity')} MW"),
        ("Formula Version", str(value.get("metadata", {}).get("formula", {}).get("version"))),
    ]


def _risk_rows(value: dict) -> list[tuple[str, str]]:
    rows = []
    for metric in value.values():
        rows.append((metric.get("name", metric.get("metric_id", "Risk")), str(metric.get("value"))))
    return rows or [("Status", "No mapped data")]


def _conversation_visualizations(dataset: dict) -> list[dict]:
    return [
        item for item in (dataset.get("conversation_visualizations") or [])
        if isinstance(item, dict) and isinstance(item.get("spec"), dict)
    ]


def _conversation_pdf_flowables(dataset: dict, styles, *, landscape_layout: bool = False) -> list:
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, PageBreak, Paragraph, Spacer, Table, TableStyle

    visualizations = _conversation_visualizations(dataset)
    if not visualizations:
        return []
    width = 245 * mm if landscape_layout else 170 * mm
    height = 128 * mm if landscape_layout else 86 * mm
    flowables = [PageBreak(), Paragraph("Charts selected from this conversation", styles["Heading1"])]
    current_section = None
    for visualization in visualizations:
        section = str(visualization.get("report_section") or "appendix").title()
        if section != current_section:
            flowables.extend([Spacer(1, 3 * mm), Paragraph(f"{section} charts", styles["Heading2"])])
            current_section = section
        title = str(visualization.get("title") or "Conversation visualization")
        summary = str(visualization.get("summary") or "")
        freshness = visualization.get("data_as_of")
        flowables.append(Paragraph(title, styles["Heading3"]))
        if summary:
            flowables.append(Paragraph(summary, styles["BodyText"]))
        if freshness:
            flowables.append(Paragraph(f"Data as of {freshness}", styles["BodyText"]))
        chart = render_visualization_spec(visualization.get("spec"))
        if chart is not None:
            flowables.extend([Spacer(1, 2 * mm), Image(chart, width=width, height=height), Spacer(1, 4 * mm)])
            continue
        rows = list(visualization.get("data_table") or [])[:20]
        columns = list(dict.fromkeys(key for row in rows for key in row))[:6]
        flowables.append(Paragraph(
            "Warning: this chart shape could not be rendered in the report; its saved data snapshot is shown instead.",
            styles["BodyText"],
        ))
        if rows and columns:
            data = [[str(column).replace("_", " ").title() for column in columns]] + [
                [str(row.get(column, ""))[:80] for column in columns] for row in rows
            ]
            flowables.append(Table(data, repeatRows=1, style=TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E2E8F0")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("PADDING", (0, 0), (-1, -1), 3),
            ])))
    return flowables


def _append_conversation_docx(document, dataset: dict, *, landscape_layout: bool = False) -> None:
    from docx.shared import Inches

    visualizations = _conversation_visualizations(dataset)
    if not visualizations:
        return
    document.add_page_break()
    document.add_heading("Charts selected from this conversation", level=1)
    current_section = None
    for visualization in visualizations:
        section = str(visualization.get("report_section") or "appendix").title()
        if section != current_section:
            document.add_heading(f"{section} charts", level=2)
            current_section = section
        document.add_heading(str(visualization.get("title") or "Conversation visualization"), level=3)
        if visualization.get("summary"):
            document.add_paragraph(str(visualization["summary"]))
        if visualization.get("data_as_of"):
            document.add_paragraph(f"Data as of {visualization['data_as_of']}")
        chart = render_visualization_spec(visualization.get("spec"))
        if chart is not None:
            document.add_picture(chart, width=Inches(8.7 if landscape_layout else 6.55))
            continue
        document.add_paragraph(
            "Warning: this chart shape could not be rendered in the report; its saved data snapshot is shown instead."
        )
        rows = list(visualization.get("data_table") or [])[:20]
        columns = list(dict.fromkeys(key for row in rows for key in row))[:6]
        if rows and columns:
            table = document.add_table(rows=1, cols=len(columns))
            table.style = "Table Grid"
            for cell, column in zip(table.rows[0].cells, columns):
                cell.text = str(column).replace("_", " ").title()
            for row in rows:
                cells = table.add_row().cells
                for cell, column in zip(cells, columns):
                    cell.text = str(row.get(column, ""))[:500]


def render_project_progress_pdf(dataset: dict, path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    story = [
        Paragraph("AKASHA", styles["Title"]),
        Paragraph("Project Progress Report", styles["Heading1"]),
        Paragraph(dataset["metadata"]["project_name"], styles["Heading2"]),
        Spacer(1, 6 * mm),
        Paragraph("Executive Summary", styles["Heading2"]),
        Paragraph(dataset["executive_summary"], styles["BodyText"]),
        Spacer(1, 5 * mm),
    ]
    chart_images = _project_report_charts(dataset)
    if chart_images:
        story.extend([Paragraph("Visual Summary", styles["Heading2"]), Spacer(1, 2 * mm)])
        for chart_image in chart_images:
            story.extend([Image(chart_image, width=170 * mm, height=86 * mm), Spacer(1, 4 * mm)])
    story.append(Paragraph("Project and Schedule", styles["Heading2"]))
    table = Table([["Metric", "Value"], *_rows(dataset)], colWidths=[70 * mm, 100 * mm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f3a5f")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.extend([table, Spacer(1, 5 * mm), Paragraph("Source Coverage", styles["Heading2"])])
    freshness = dataset["metadata"]["source_freshness"]
    story.append(Table([["Source", "Last synchronized"], *[
        [name, str(value or "No mapped data")] for name, value in freshness.items()
    ]], colWidths=[45 * mm, 125 * mm], style=TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
        ("PADDING", (0, 0), (-1, -1), 6),
    ])))
    for heading, key in [
        ("SAP Procurement", "procurement"),
        ("TC Transmission", "transmission"),
        ("Pulse Quality", "quality"),
        ("Capacity Milestones", "capacity"),
        ("Named Risk Metrics", "risk"),
    ]:
        story.extend([Spacer(1, 5 * mm), Paragraph(heading, styles["Heading2"])])
        value = dataset.get(key, {})
        rows = _capacity_rows(value) if key == "capacity" else _risk_rows(value) if key == "risk" else _domain_rows(value)
        story.append(Table([["Metric", "Value"], *rows], colWidths=[70 * mm, 100 * mm], style=TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
            ("PADDING", (0, 0), (-1, -1), 6),
        ])))
    # Give the activity register a clean page boundary instead of leaving a
    # handful of orphaned rows after the domain metric sections.
    story.extend([PageBreak(), Paragraph("In-Progress Activities", styles["Heading2"])])
    activities = dataset["in_progress_activities"].get("activities") or []
    activity_rows = [[row.get("activity_id"), row.get("name"), f"{row.get('percent_complete')}%"] for row in activities]
    story.append(Table(
        [["ID", "Activity", "Complete"], *activity_rows] if activity_rows else [["Status"], ["No in-progress activities"]],
        repeatRows=1,
        colWidths=[30 * mm, 105 * mm, 30 * mm] if activity_rows else None,
        style=TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("PADDING", (0, 0), (-1, -1), 4),
    ])))
    story.extend(_conversation_pdf_flowables(dataset, styles))
    SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm).build(story)


def render_project_progress_docx(dataset: dict, path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, RGBColor

    document = Document()
    _configure_docx(document, running_label="Akasha | Project Progress Report")
    title = document.add_heading("AKASHA", 0)
    title.runs[0].font.color.rgb = RGBColor(31, 58, 95)
    document.add_heading("Project Progress Report", level=1)
    document.add_heading(dataset["metadata"]["project_name"], level=2)
    document.add_heading("Executive Summary", level=2)
    document.add_paragraph(dataset["executive_summary"])
    chart_images = _project_report_charts(dataset)
    if chart_images:
        document.add_heading("Visual Summary", level=2)
        for chart_image in chart_images:
            document.add_picture(chart_image, width=Inches(6.55))
    document.add_heading("Project and Schedule", level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    for label, value in _rows(dataset):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    _style_docx_table(table, [1.875, 4.625])
    document.add_heading("Source Coverage", level=2)
    for name, value in dataset["metadata"]["source_freshness"].items():
        document.add_paragraph(f"{name}: {value or 'No mapped data'}")
    for heading, key in [
        ("SAP Procurement", "procurement"),
        ("TC Transmission", "transmission"),
        ("Pulse Quality", "quality"),
        ("Capacity Milestones", "capacity"),
        ("Named Risk Metrics", "risk"),
    ]:
        document.add_heading(heading, level=2)
        domain_table = document.add_table(rows=1, cols=2)
        domain_table.style = "Table Grid"
        domain_table.rows[0].cells[0].text = "Metric"
        domain_table.rows[0].cells[1].text = "Value"
        value = dataset.get(key, {})
        rows = _capacity_rows(value) if key == "capacity" else _risk_rows(value) if key == "risk" else _domain_rows(value)
        for label, value in rows:
            cells = domain_table.add_row().cells
            cells[0].text = label
            cells[1].text = value
        _style_docx_table(domain_table, [1.875, 4.625])
    document.add_heading("In-Progress Activities", level=2)
    activity_table = document.add_table(rows=1, cols=3)
    activity_table.style = "Table Grid"
    for cell, value in zip(activity_table.rows[0].cells, ["ID", "Activity", "Complete"]):
        cell.text = value
    for row in dataset["in_progress_activities"].get("activities") or []:
        cells = activity_table.add_row().cells
        cells[0].text = str(row.get("activity_id") or "")
        cells[1].text = str(row.get("name") or "")
        cells[2].text = f"{row.get('percent_complete')}%"
    _style_docx_table(activity_table, [1.2, 4.1, 1.2])
    _append_conversation_docx(document, dataset)
    document.add_paragraph("Confidential - generated by Akasha from synchronized source data.")
    document.save(str(path))


def _portfolio_project_rows(dataset: dict) -> list[list[str]]:
    return [[
        str(row.get("project_name") or row.get("project_id")),
        f"{float(row.get('progress_pct')):.1f}%" if row.get("progress_pct") is not None else "Unavailable",
        str(row.get("status") or "Unavailable"),
        str(row.get("forecast_finish") or "Unavailable"),
        str(row.get("finish_date_variance_days") if row.get("finish_date_variance_days") is not None else "Unavailable"),
    ] for row in dataset.get("projects", [])]


def _portfolio_report_charts(dataset: dict) -> list[BytesIO]:
    visualizations = dataset.get("report_visualizations") or {}
    images = []
    progress_spec = visualizations.get("project_progress") or {}
    if isinstance(progress_spec, dict) and progress_spec.get("schema_version") == "visualization.v1":
        progress_chart = render_visualization_spec(progress_spec)
    else:
        progress = [
            (row["project_name"], float(row["progress_pct"]))
            for row in progress_spec
            if row.get("progress_pct") is not None
        ]
        progress.sort(key=lambda row: row[1], reverse=True)
        progress_chart = _horizontal_bar_chart(
            progress,
            "Portfolio Progress Comparison",
            "Top projects by authoritative current P6 progress",
        )
    if progress_chart:
        images.append(progress_chart)
    status_spec = visualizations.get("schedule_status") or {}
    status_chart = (
        render_visualization_spec(status_spec)
        if isinstance(status_spec, dict) and status_spec.get("schema_version") == "visualization.v1"
        else _status_donut_chart(
            status_spec,
            "Portfolio Schedule Status",
            "Current status distribution at the latest synchronized cutoff",
        )
    )
    if status_chart:
        images.append(status_chart)
    return images


def render_portfolio_progress_pdf(dataset: dict, path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    metadata = dataset["metadata"]
    summary = dataset["summary"]
    story = [
        Paragraph("AKASHA", styles["Title"]),
        Paragraph("Portfolio Progress Report", styles["Heading1"]),
        Paragraph(metadata["portfolio"], styles["Heading2"]),
        Paragraph(
            f"Current period: {metadata.get('period_start') or 'Unavailable'} through "
            f"{metadata.get('reporting_cutoff') or 'latest synchronized cutoff'}",
            styles["BodyText"],
        ),
        Spacer(1, 4 * mm),
        Paragraph("Executive Summary", styles["Heading2"]),
        Paragraph(dataset["executive_summary"], styles["BodyText"]),
        Spacer(1, 4 * mm),
    ]
    kpis = [
        ["Projects", "P6 available", "Delayed", "On track", "Completed", "P6 unavailable"],
        [summary["total_projects"], summary["projects_with_p6"], summary["delayed"], summary["on_track"], summary["completed"], summary["p6_unavailable"]],
    ]
    story.append(Table(kpis, repeatRows=1, style=TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1769AA")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ("PADDING", (0, 0), (-1, -1), 7),
    ])))
    chart_images = _portfolio_report_charts(dataset)
    if chart_images:
        chart_cells = [
            Image(chart_image, width=120 * mm, height=64 * mm)
            for chart_image in chart_images[:2]
        ]
        story.extend([
            Spacer(1, 5 * mm),
            Table([chart_cells], colWidths=[123 * mm] * len(chart_cells), style=TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 1),
                ("RIGHTPADDING", (0, 0), (-1, -1), 1),
            ])),
        ])
    story.extend([PageBreak(), Paragraph("Project Detail", styles["Heading2"])])
    detail = Table(
        [["Project", "Progress", "Schedule", "Forecast finish", "Variance (days)"], *_portfolio_project_rows(dataset)],
        repeatRows=1,
        colWidths=[88 * mm, 28 * mm, 32 * mm, 38 * mm, 34 * mm],
        style=TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1769AA")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
            ("PADDING", (0, 0), (-1, -1), 4),
        ]),
    )
    story.append(detail)
    for limitation in metadata.get("limitations") or []:
        story.extend([Spacer(1, 3 * mm), Paragraph(f"Limitation: {limitation}", styles["BodyText"])])
    story.extend(_conversation_pdf_flowables(dataset, styles, landscape_layout=True))
    SimpleDocTemplate(
        str(path), pagesize=landscape(A4), rightMargin=14 * mm, leftMargin=14 * mm,
        topMargin=12 * mm, bottomMargin=12 * mm,
    ).build(story)


def render_portfolio_progress_docx(dataset: dict, path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, RGBColor

    document = Document()
    _configure_docx(
        document,
        landscape=True,
        running_label="Akasha | Portfolio Progress Report",
    )
    title = document.add_heading("AKASHA", 0)
    title.runs[0].font.color.rgb = RGBColor(23, 105, 170)
    document.add_heading("Portfolio Progress Report", level=1)
    document.add_heading(dataset["metadata"]["portfolio"], level=2)
    document.add_paragraph(dataset["executive_summary"])
    document.add_heading("Visual Summary", level=2)
    chart_images = _portfolio_report_charts(dataset)
    if chart_images:
        visual_table = document.add_table(rows=1, cols=min(2, len(chart_images)))
        for cell, chart_image in zip(visual_table.rows[0].cells, chart_images[:2]):
            cell.width = Inches(4.45)
            cell.paragraphs[0].add_run().add_picture(chart_image, width=Inches(4.25))
    document.add_heading("Portfolio KPI Summary", level=2)
    summary = dataset["summary"]
    for label, key in [
        ("Total projects", "total_projects"), ("P6 available", "projects_with_p6"),
        ("Delayed", "delayed"), ("On track", "on_track"),
        ("Completed", "completed"), ("P6 unavailable", "p6_unavailable"),
    ]:
        document.add_paragraph(f"{label}: {summary[key]}")
    document.add_heading("Project Detail", level=2)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ["Project", "Progress", "Schedule", "Forecast finish", "Variance (days)"]):
        cell.text = value
    for row in _portfolio_project_rows(dataset):
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    _style_docx_table(table, [3.5, 1.0, 1.2, 1.7, 1.35])
    document.add_heading("Limitations", level=2)
    for limitation in dataset["metadata"].get("limitations") or []:
        document.add_paragraph(limitation, style="List Bullet")
    _append_conversation_docx(document, dataset, landscape_layout=True)
    document.add_paragraph("Confidential - generated by Akasha from synchronized source data.")
    document.save(str(path))


def _comparison_report_charts(dataset: dict) -> list[BytesIO]:
    return [
        image for image in (
            render_visualization_spec(spec)
            for spec in (dataset.get("report_visualizations") or {}).values()
        ) if image is not None
    ][:4]


def _comparison_metric_rows(dataset: dict) -> list[list[str]]:
    projects = dataset.get("projects") or []
    metrics = [
        ("Capacity", lambda row: f"{row.get('capacity_mwac')} MW" if row.get("capacity_mwac") is not None else "Unavailable"),
        ("SPV", lambda row: str(row.get("spv_name") or "Unavailable")),
        ("Progress", lambda row: f"{float(row.get('progress_pct')):.1f}%" if row.get("progress_pct") is not None else "Unavailable"),
        ("Status", lambda row: str(row.get("status") or "Unavailable")),
        ("Activities complete", lambda row: str(row.get("completed_activities") or 0)),
        ("Activities in progress", lambda row: str(row.get("in_progress_activities") or 0)),
        ("Activities not started", lambda row: str(row.get("not_started_activities") or 0)),
        ("Planned duration", lambda row: f"{float(row.get('planned_duration')):.0f} hrs" if row.get("planned_duration") is not None else "Unavailable"),
        ("Actual duration", lambda row: f"{float(row.get('actual_duration')):.0f} hrs" if row.get("actual_duration") is not None else "Unavailable"),
        ("Remaining duration", lambda row: f"{float(row.get('remaining_duration')):.0f} hrs" if row.get("remaining_duration") is not None else "Unavailable"),
        ("Baseline finish", lambda row: str(row.get("baseline_finish") or "Unavailable")),
        ("Forecast finish", lambda row: str(row.get("forecast_finish") or "Unavailable")),
        ("Baseline slip", lambda row: f"{int(row.get('baseline_slip_days'))} days" if row.get("baseline_slip_days") is not None else "Unavailable"),
        ("Data as of", lambda row: str(row.get("data_as_of") or "Unavailable")),
    ]
    return [[label, *[formatter(project) for project in projects]] for label, formatter in metrics]


def render_project_comparison_pdf(dataset: dict, path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    metadata = dataset["metadata"]
    projects = dataset["projects"]
    story = [
        Paragraph("AKASHA", styles["Title"]),
        Paragraph("Project Comparison Report", styles["Heading1"]),
        Paragraph(" vs ".join(metadata["project_names"]), styles["Heading2"]),
        Paragraph(f"Latest reporting cutoff: {metadata.get('reporting_cutoff') or 'Unavailable'}", styles["BodyText"]),
        Spacer(1, 4 * mm),
        Paragraph("Executive Summary", styles["Heading2"]),
        Paragraph(dataset["executive_summary"], styles["BodyText"]),
        Spacer(1, 4 * mm),
        Paragraph("Visual Comparison", styles["Heading2"]),
    ]
    images = _comparison_report_charts(dataset)
    for offset in range(0, len(images), 2):
        cells = [Image(image, width=120 * mm, height=64 * mm) for image in images[offset:offset + 2]]
        story.append(Table([cells], colWidths=[123 * mm] * len(cells), style=TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 1),
            ("RIGHTPADDING", (0, 0), (-1, -1), 1),
        ])))
        story.append(Spacer(1, 3 * mm))
    story.extend([PageBreak(), Paragraph("Key Metrics Comparison", styles["Heading2"])])
    table_data = [["Metric", *[project["project_name"] for project in projects]], *_comparison_metric_rows(dataset)]
    available_width = 255 * mm
    metric_width = 42 * mm
    project_width = (available_width - metric_width) / len(projects)
    story.append(Table(table_data, repeatRows=1, colWidths=[metric_width, *([project_width] * len(projects))], style=TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1769AA")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("PADDING", (0, 0), (-1, -1), 5),
    ])))
    story.extend([Spacer(1, 4 * mm), Paragraph("Limitations", styles["Heading2"])])
    for limitation in metadata.get("limitations") or []:
        story.append(Paragraph(f"• {limitation}", styles["BodyText"]))
    story.extend(_conversation_pdf_flowables(dataset, styles, landscape_layout=True))
    SimpleDocTemplate(
        str(path), pagesize=landscape(A4), rightMargin=14 * mm, leftMargin=14 * mm,
        topMargin=12 * mm, bottomMargin=12 * mm,
    ).build(story)


def render_project_comparison_docx(dataset: dict, path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, RGBColor

    document = Document()
    _configure_docx(document, landscape=True, running_label="Akasha | Project Comparison Report")
    title = document.add_heading("AKASHA", 0)
    title.runs[0].font.color.rgb = RGBColor(23, 105, 170)
    document.add_heading("Project Comparison Report", level=1)
    document.add_heading(" vs ".join(dataset["metadata"]["project_names"]), level=2)
    document.add_paragraph(dataset["executive_summary"])
    document.add_heading("Visual Comparison", level=2)
    images = _comparison_report_charts(dataset)
    for offset in range(0, len(images), 2):
        row_table = document.add_table(rows=1, cols=min(2, len(images) - offset))
        for cell, chart_image in zip(row_table.rows[0].cells, images[offset:offset + 2]):
            cell.width = Inches(4.45)
            cell.paragraphs[0].add_run().add_picture(chart_image, width=Inches(4.25))
    document.add_heading("Key Metrics Comparison", level=2)
    projects = dataset["projects"]
    table = document.add_table(rows=1, cols=1 + len(projects))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ["Metric", *[row["project_name"] for row in projects]]):
        cell.text = value
    for row in _comparison_metric_rows(dataset):
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    _style_docx_table(table, [1.5, *([3.8] * len(projects))])
    document.add_heading("Limitations", level=2)
    for limitation in dataset["metadata"].get("limitations") or []:
        document.add_paragraph(limitation, style="List Bullet")
    _append_conversation_docx(document, dataset, landscape_layout=True)
    document.add_paragraph("Confidential - generated by Akasha from synchronized source data.")
    document.save(str(path))
