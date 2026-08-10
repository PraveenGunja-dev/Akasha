"""
Akasha Platform — Report Export Service (Word .docx & PDF .pdf)

Generates branded Adani Renewables executive reports from chatbot Markdown responses
and metadata, featuring:
- Adani Renewables Header & Branding
- Watermark ("ADANI RENEWABLES — CONFIDENTIAL")
- Styled Headings, Bullet Points, and Tables
- Metadata summary footer (Sources, Project IDs, Execution Latency)
"""

import io
import os
import re
import logging
from datetime import datetime
from typing import Dict, Any, Optional

# --- Word Generation ---
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# --- PDF Generation ---
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
)
from reportlab.pdfgen import canvas

logger = logging.getLogger(__name__)

# --- Colors ---
ADANI_NAVY = RGBColor(11, 116, 176)      # #0B74B0
ADANI_PURPLE = RGBColor(117, 71, 156)    # #75479C
ADANI_DARK = RGBColor(15, 23, 42)        # #0F172A
ADANI_MUTED = RGBColor(100, 116, 139)    # #64748B

RL_NAVY = colors.HexColor("#0B74B0")
RL_PURPLE = colors.HexColor("#75479C")
RL_DARK = colors.HexColor("#0F172A")
RL_MUTED = colors.HexColor("#64748B")
RL_BG_LIGHT = colors.HexColor("#F8FAFC")
RL_BORDER = colors.HexColor("#E2E8F0")


# ==============================================================================
# 1. WORD (.DOCX) GENERATOR
# ==============================================================================

def set_cell_background(cell, fill_hex: str):
    """Set background color of a Word table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)


def add_docx_watermark(doc, text="ADANI RENEWABLES — CONFIDENTIAL"):
    """Adds a subtle header watermark to Word document."""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"🔒 {text}")
    run.font.size = Pt(8.5)
    run.font.color.rgb = ADANI_MUTED
    run.font.name = "Calibri"


def parse_markdown_table(table_text: str):
    """Parses markdown table into header list and row lists."""
    lines = [line.strip() for line in table_text.strip().split("\n") if line.strip()]
    if len(lines) < 2:
        return None, []
    
    # Filter out separator line (|---|---|)
    cleaned_lines = [l for l in lines if not re.match(r'^\|?\s*:?-+:?\s*(\|?\s*:?-+:?\s*)*\|?$', l)]
    
    table_data = []
    for line in cleaned_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        table_data.append(cells)
        
    if not table_data:
        return None, []
        
    headers = table_data[0]
    rows = table_data[1:]
    return headers, rows


def _clean_ascii(text: str) -> str:
    """Removes non-latin/non-ASCII characters that default PIL font cannot render."""
    if not text:
        return ""
    text = str(text).replace("—", "-").replace("–", "-").replace("📊", "").replace("📈", "").replace("🔒", "").replace("📋", "").strip()
    return re.sub(r'[^\x00-\x7F]+', ' ', text)


def _render_chart_spec_to_png(viz: dict) -> io.BytesIO:
    """Renders an ECharts spec dict into a crisp, 300 DPI executive chart image using Matplotlib."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import numpy as np

        raw_title = viz.get("title") or viz.get("chart_type", "Executive Visualization")
        title = _clean_ascii(raw_title)
        spec = viz.get("spec") or viz.get("option") or {}
        chart_type = viz.get("chart_type", "").lower()

        series_list = spec.get("series", [])
        x_axis = spec.get("xAxis", {})
        y_axis = spec.get("yAxis", {})

        categories = []
        if isinstance(y_axis, dict) and y_axis.get("data"):
            categories = y_axis.get("data", [])
        elif isinstance(x_axis, dict) and x_axis.get("data"):
            categories = x_axis.get("data", [])
        elif isinstance(y_axis, list) and len(y_axis) > 0 and isinstance(y_axis[0], dict) and y_axis[0].get("data"):
            categories = y_axis[0].get("data", [])
        elif isinstance(x_axis, list) and len(x_axis) > 0 and isinstance(x_axis[0], dict) and x_axis[0].get("data"):
            categories = x_axis[0].get("data", [])

        categories = [_clean_ascii(c.get("value", c) if isinstance(c, dict) else c)[:35] for c in categories]
        colors_palette = ["#2563EB", "#059669", "#D97706", "#7C3AED", "#0891B2", "#DC2626"]

        # 1. MULTI-SERIES GROUPED BAR (e.g. EVM Health SPI vs CPI, Vendor Performance)
        if len(series_list) > 1 and (series_list[0].get("type") in ("bar", None) or "evm" in chart_type):
            fig, ax = plt.subplots(figsize=(9.5, 4.5), dpi=300)
            fig.patch.set_facecolor('#FFFFFF')
            ax.set_facecolor('#F8FAFC')

            ax.set_title(title, fontsize=12, fontweight='bold', color='#0F172A', pad=15, loc='left')

            n_cats = len(categories) if categories else 1
            y = np.arange(n_cats)
            n_series = len(series_list)
            total_height = 0.55
            bar_height = total_height / n_series

            all_vals = []
            for i, s in enumerate(series_list):
                s_name = s.get("name", f"Series {i+1}")
                s_data = s.get("data", [])
                vals = [d.get("value", d) if isinstance(d, dict) else d for d in s_data]
                vals = [v if isinstance(v, (int, float)) else 0 for v in vals]
                all_vals.extend(vals)

                s_color = s.get("itemStyle", {}).get("color") or colors_palette[i % len(colors_palette)]
                pos = y + (i - n_series/2 + 0.5) * bar_height
                rects = ax.barh(pos, vals, bar_height, label=s_name, color=s_color, edgecolor='none')

                for bar in rects:
                    w = bar.get_width()
                    if w > 0:
                        ax.text(w + 0.02, bar.get_y() + bar.get_height()/2, f"{w:.2f}" if isinstance(w, float) else f"{w}", ha='left', va='center', fontsize=8.5, fontweight='bold', color='#1E293B')

            if categories:
                ax.set_yticks(y)
                ax.set_yticklabels(categories, fontsize=9.5, fontweight='bold', color='#0F172A')

            max_v = max(all_vals) if all_vals else 1
            ax.set_xlim(0, max_v * 1.18)

            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_color('#CBD5E1')
            ax.spines['bottom'].set_color('#CBD5E1')
            ax.xaxis.grid(True, linestyle='--', alpha=0.5, color='#CBD5E1')
            ax.legend(loc='upper right', frameon=True, facecolor='#FFFFFF', edgecolor='#E2E8F0', fontsize=9)
            plt.tight_layout()

        # 2. DONUT / PIE CHARTS (Risk Matrix, Activity Status, Transmission Status)
        elif ("matrix" in chart_type or "status" in chart_type or "donut" in chart_type or (series_list and series_list[0].get("type") == "pie")):
            fig, ax = plt.subplots(figsize=(8.0, 4.5), dpi=300)
            fig.patch.set_facecolor('#FFFFFF')
            ax.set_facecolor('#FFFFFF')

            ax.set_title(title, fontsize=12, fontweight='bold', color='#0F172A', pad=15, loc='left')

            pie_data = series_list[0].get("data", []) if series_list else []
            labels = []
            values = []
            colors = []

            for i, item in enumerate(pie_data):
                if isinstance(item, dict):
                    raw_lbl = _clean_ascii(item.get("name", f"Item {i+1}"))
                    lbl = raw_lbl.split(":")[0].strip() if ":" in raw_lbl else raw_lbl
                    val = item.get("value", 0)
                    clr = item.get("itemStyle", {}).get("color") or colors_palette[i % len(colors_palette)]
                else:
                    raw_lbl = categories[i] if i < len(categories) else f"Item {i+1}"
                    lbl = raw_lbl.split(":")[0].strip() if ":" in raw_lbl else raw_lbl
                    val = item
                    clr = colors_palette[i % len(colors_palette)]
                labels.append(lbl)
                values.append(val if isinstance(val, (int, float)) else 0)
                colors.append(clr)

            total = sum(values) if values else 1

            wedges, texts, autotexts = ax.pie(
                values,
                labels=None,
                autopct=lambda pct: f'{pct:.1f}%' if pct > 4 else '',
                startangle=90,
                pctdistance=0.75,
                colors=colors,
                wedgeprops=dict(width=0.4, edgecolor='#FFFFFF', linewidth=2)
            )
            plt.setp(autotexts, size=8.5, weight="bold", color="white")

            ax.text(0, 0, f"Total\n{total}", ha='center', va='center', fontsize=11, fontweight='bold', color='#0F172A')

            legend_labels = [f"{lbl}: {val}" for lbl, val in zip(labels, values)]
            ax.legend(wedges, legend_labels, loc="center left", bbox_to_anchor=(1, 0, 0.5, 1), frameon=True, facecolor='#F8FAFC', edgecolor='#E2E8F0', fontsize=9)
            plt.tight_layout()

        # 3. SINGLE-SERIES HORIZONTAL BAR (Project Comparison, Delayed Activities)
        else:
            fig, ax = plt.subplots(figsize=(9.5, 4.5), dpi=300)
            fig.patch.set_facecolor('#FFFFFF')
            ax.set_facecolor('#F8FAFC')

            ax.set_title(title, fontsize=12, fontweight='bold', color='#0F172A', pad=15, loc='left')

            series = series_list[0] if series_list else {}
            raw_data = series.get("data", [])

            vals = []
            bar_colors = []
            default_color = series.get("itemStyle", {}).get("color") or "#0B74B0"

            for i, d in enumerate(raw_data):
                if isinstance(d, dict):
                    vals.append(d.get("value", 0))
                    bar_colors.append(d.get("itemStyle", {}).get("color") or default_color)
                else:
                    vals.append(d)
                    bar_colors.append(default_color)

            vals = [v if isinstance(v, (int, float)) else 0 for v in vals]
            n_cats = len(categories) if categories else len(vals)
            y = np.arange(n_cats)

            rects = ax.barh(y, vals, height=0.45, color=bar_colors, edgecolor='none')

            if categories:
                ax.set_yticks(y)
                ax.set_yticklabels(categories, fontsize=9.5, fontweight='bold', color='#0F172A')

            max_v = max(vals) if vals else 1
            ax.set_xlim(0, max_v * 1.18)

            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_color('#CBD5E1')
            ax.spines['bottom'].set_color('#CBD5E1')
            ax.xaxis.grid(True, linestyle='--', alpha=0.5, color='#CBD5E1')

            for bar in rects:
                w = bar.get_width()
                if w > 0:
                    ax.text(w + (max_v * 0.015), bar.get_y() + bar.get_height()/2, f"{w:.1f}%" if "%" in title else f"{w}", ha='left', va='center', fontsize=8.5, fontweight='bold', color='#0F172A')

            plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception as e:
        logger.error(f"Matplotlib chart rendering error: {e}")
        return None

        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
    except Exception as e:
        logger.error(f"Failed to render PIL chart PNG: {e}")
        return None


def generate_docx_report(title: str, content: str, metadata: Optional[Dict[str, Any]] = None, images: Optional[list] = None, visualizations: Optional[list] = None) -> bytes:
    """
    Generates a professionally styled Word (.docx) document from Markdown response and chart images/visualizations.
    Returns bytes buffer.
    """
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    add_docx_watermark(doc)
    
    # --- Adani Logo & Title Banner ---
    clean_logo_path = os.path.join(os.path.dirname(__file__), "..", "assets", "adani_clean.png")
    fallback_logo_path = os.path.join(os.path.dirname(__file__), "..", "assets", "adani.png")
    logo_path = clean_logo_path if os.path.exists(clean_logo_path) else fallback_logo_path
    
    if os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Inches(1.2))
        except Exception:
            pass

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(4)
    run_brand = title_p.add_run("ADANI RENEWABLES — AKASHA AI\n")
    run_brand.font.name = "Calibri"
    run_brand.font.size = Pt(10)
    run_brand.font.bold = True
    run_brand.font.color.rgb = ADANI_PURPLE
    
    run_title = title_p.add_run(title or "Executive Project Intelligence Report")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = ADANI_NAVY

    # Sub-header Meta
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(16)
    date_str = datetime.now().strftime("%B %d, %Y | %H:%M IST")
    sub_run = sub_p.add_run(f"Generated on {date_str}  •  Classification: CONFIDENTIAL")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = ADANI_MUTED
    
    # Horizontal Divider Line
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_after = Pt(14)
    p_hr_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="0B74B0"/></w:pBdr>')
    p_hr._p.get_or_add_pPr().append(p_hr_border)

    # --- Content Parsing ---
    lines = content.split("\n")
    in_table = False
    table_buffer = []

    def flush_table_buffer():
        nonlocal table_buffer
        if not table_buffer:
            return
        headers, rows = parse_markdown_table("\n".join(table_buffer))
        table_buffer = []
        if not headers:
            return
            
        num_cols = len(headers)
        tbl = doc.add_table(rows=1 + len(rows), cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        
        # Header Row
        hdr_cells = tbl.rows[0].cells
        for idx, text in enumerate(headers):
            hdr_cells[idx].text = text
            set_cell_background(hdr_cells[idx], "0B74B0")
            set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
            for p in hdr_cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(10)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    
        # Data Rows
        for r_idx, row_data in enumerate(rows):
            row_cells = tbl.rows[r_idx + 1].cells
            bg_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx in range(min(num_cols, len(row_data))):
                row_cells[c_idx].text = row_data[c_idx]
                set_cell_background(row_cells[c_idx], bg_color)
                set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
                for p in row_cells[c_idx].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        r.font.name = "Calibri"
                        r.font.size = Pt(9.5)
                        r.font.color.rgb = ADANI_DARK

        p_spacer = doc.add_paragraph()
        p_spacer.paragraph_format.space_after = Pt(10)

    for line in lines:
        stripped = line.strip()
        
        # Check table
        if "|" in stripped:
            in_table = True
            table_buffer.append(line)
            continue
        elif in_table:
            in_table = False
            flush_table_buffer()
            
        if not stripped:
            continue
            
        # Headings
        if stripped.startswith("# "):
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            r = h.add_run(stripped[2:])
            r.font.name = "Calibri"
            r.font.size = Pt(15)
            r.font.bold = True
            r.font.color.rgb = ADANI_NAVY
        elif stripped.startswith("## "):
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            r = h.add_run(stripped[3:])
            r.font.name = "Calibri"
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = ADANI_PURPLE
        elif stripped.startswith("### "):
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            r = h.add_run(stripped[4:])
            r.font.name = "Calibri"
            r.font.size = Pt(11.5)
            r.font.bold = True
            r.font.color.rgb = ADANI_DARK
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            # Inline bold formatting
            parts = re.split(r'(\*\*.*?\*\*)', stripped[2:])
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = "Calibri"
                r.font.size = Pt(10.5)
                r.font.color.rgb = ADANI_DARK
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = "Calibri"
                r.font.size = Pt(10.5)
                r.font.color.rgb = ADANI_DARK

    if in_table:
        flush_table_buffer()

    # --- Footer Metadata ---
    if metadata:
        doc.add_paragraph().paragraph_format.space_before = Pt(16)
        meta_table = doc.add_table(rows=1, cols=1)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = meta_table.rows[0].cells[0]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        mp = cell.paragraphs[0]
        mp.paragraph_format.space_after = Pt(2)
        r_meta = mp.add_run("📋 Executive Summary & Execution Audit\n")
        r_meta.font.name = "Calibri"
        r_meta.font.bold = True
        r_meta.font.size = Pt(10)
        r_meta.font.color.rgb = ADANI_NAVY
        
        sources = ", ".join(metadata.get("sources", [])) or "PostgreSQL Database Engine"
        projects = ", ".join(metadata.get("project_ids", [])) or "Portfolio Level"
        latency = metadata.get("latency_ms", "N/A")
        
        info_str = f"• Projects Referenced: {projects}\n• Primary Data Sources: {sources}\n• Query Latency: {latency} ms"
        r_info = mp.add_run(info_str)
        r_info.font.name = "Calibri"
        r_info.font.size = Pt(9)
        r_info.font.color.rgb = ADANI_MUTED

    # --- Add Visualizations / Chart Images ---
    chart_bytes_list: List[bytes] = []
    
    # Priority 1: Use high-res browser Base64 screenshots (exact WYSIWYG match to chatbot)
    if images and any(images):
        import base64
        for img_str in images:
            if not img_str:
                continue
            try:
                clean_b64 = img_str.split(",")[-1] if "," in str(img_str) else str(img_str)
                img_bytes = base64.b64decode(clean_b64)
                if len(img_bytes) > 200:
                    chart_bytes_list.append(img_bytes)
            except Exception as e:
                logger.warning(f"Failed to decode image string for Word report: {e}")

    # Priority 2: Fallback to high-DPI 300 DPI Matplotlib Engine if raw browser images absent
    if not chart_bytes_list and visualizations and len(visualizations) > 0:
        for viz in visualizations:
            try:
                cs = _render_chart_spec_to_png(viz)
                if cs:
                    chart_bytes_list.append(cs.getvalue())
            except Exception as e:
                logger.warning(f"Failed to render chart stream for Word report: {e}")

    if chart_bytes_list:
        doc.add_paragraph().paragraph_format.space_before = Pt(16)
        p_v = doc.add_paragraph()
        p_v.paragraph_format.space_after = Pt(12)
        r_v = p_v.add_run("Executive Visualizations & Key Charts")
        r_v.font.name = "Calibri"
        r_v.font.size = Pt(14)
        r_v.font.bold = True
        r_v.font.color.rgb = ADANI_NAVY

        # Render each chart full-width (6.2 inches) centered — NEVER squish in tiny table cells
        for img_b in chart_bytes_list:
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_after = Pt(14)
                stream = io.BytesIO(img_b)
                p_img.add_run().add_picture(stream, width=Inches(6.2))
            except Exception as e:
                logger.warning(f"Failed to add picture to Word report: {e}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ==============================================================================
# 2. PDF GENERATOR
# ==============================================================================

class NumberedCanvas(canvas.Canvas):
    """Canvas that handles multi-pass page numbering and Adani Watermark."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_number(self, page_count):
        self.saveState()
        # --- Confidential Watermark ---
        self.setFont("Helvetica-Bold", 32)
        self.setFillColor(colors.HexColor("#0B74B0"))
        self.setLineWidth(0)
        if hasattr(self, 'setFillAlpha'):
            self.setFillAlpha(0.04)
        
        self.translate(300, 400)
        self.rotate(40)
        self.drawCentredString(0, 0, "ADANI RENEWABLES — CONFIDENTIAL")
        self.restoreState()
        
        self.saveState()
        # --- Header Bar ---
        self.setFillColor(RL_NAVY)
        self.rect(36, 805, 540, 4, fill=1, stroke=0)
        
        self.setFont("Helvetica-Bold", 8)
        self.setFillColor(RL_PURPLE)
        self.drawString(36, 814, "ADANI RENEWABLES — AKASHA INTELLIGENCE REPORT")
        
        # --- Footer Page Numbers ---
        self.setFont("Helvetica", 8)
        self.setFillColor(RL_MUTED)
        self.drawString(36, 24, "Classification: STRICTLY CONFIDENTIAL")
        page_str = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(576, 24, page_str)
        
        self.setStrokeColor(RL_BORDER)
        self.setLineWidth(0.5)
        self.line(36, 36, 576, 36)
        
        self.restoreState()


def generate_pdf_report(title: str, content: str, metadata: Optional[Dict[str, Any]] = None, images: Optional[list] = None, visualizations: Optional[list] = None) -> bytes:
    """
    Generates a styled PDF report with Adani Header, Watermark, and Table formatting.
    Returns bytes buffer.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom Typography Styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=RL_NAVY,
        spaceAfter=4
    )
    
    sub_style = ParagraphStyle(
        'DocSub',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=RL_MUTED,
        spaceAfter=14
    )
    
    h1_style = ParagraphStyle(
        'H1',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=RL_NAVY,
        spaceBefore=14,
        spaceAfter=6
    )
    
    h2_style = ParagraphStyle(
        'H2',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        textColor=RL_PURPLE,
        spaceBefore=10,
        spaceAfter=4
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=14,
        textColor=RL_DARK,
        spaceAfter=6
    )
    
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=14,
        textColor=RL_DARK,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=3
    )
    
    story = []

    # Title & Header
    story.append(Paragraph("ADANI RENEWABLES — AKASHA INTELLIGENCE", ParagraphStyle('Brand', fontName='Helvetica-Bold', fontSize=9, textColor=RL_PURPLE, spaceAfter=2)))
    story.append(Paragraph(title or "Executive Project Intelligence Report", title_style))
    date_str = datetime.now().strftime("%B %d, %Y | %H:%M IST")
    story.append(Paragraph(f"Generated on {date_str}  •  Classification: CONFIDENTIAL", sub_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=RL_NAVY, spaceAfter=14))

    # --- Body Content & Table Parsing ---
    lines = content.split("\n")
    in_table = False
    table_buffer = []

    def flush_pdf_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        headers, rows = parse_markdown_table("\n".join(table_buffer))
        table_buffer = []
        if not headers:
            return
            
        data = [[Paragraph(f"<b>{h}</b>", ParagraphStyle('TH', fontName='Helvetica-Bold', fontSize=9, textColor=colors.white)) for h in headers]]
        for row in rows:
            data.append([Paragraph(r, body_style) for r in row])
            
        col_w = [520 / len(headers)] * len(headers)
        t = Table(data, colWidths=col_w)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), RL_NAVY),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [RL_BG_LIGHT, colors.white]),
            ('GRID', (0, 0), (-1, -1), 0.5, RL_BORDER),
        ]))
        story.append(t)
        story.append(Spacer(1, 10))

    for line in lines:
        stripped = line.strip()
        if "|" in stripped:
            in_table = True
            table_buffer.append(line)
            continue
        elif in_table:
            in_table = False
            flush_pdf_table()
            
        if not stripped:
            continue
            
        if stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], h1_style))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], h2_style))
        elif stripped.startswith("### "):
            story.append(Paragraph(stripped[4:], h2_style))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            formatted = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', stripped[2:])
            story.append(Paragraph(f"• {formatted}", bullet_style))
        else:
            formatted = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
            story.append(Paragraph(formatted, body_style))

    if in_table:
        flush_pdf_table()

    # --- Add Visualizations / Chart Images to PDF ---
    chart_bytes_list: List[bytes] = []
    from reportlab.platypus import Image as RLImage
    
    # Priority 1: Use high-res browser Base64 screenshots (exact WYSIWYG match to chatbot)
    if images and any(images):
        import base64
        for img_str in images:
            if not img_str:
                continue
            try:
                clean_b64 = img_str.split(",")[-1] if "," in str(img_str) else str(img_str)
                img_bytes = base64.b64decode(clean_b64)
                if len(img_bytes) > 200:
                    chart_bytes_list.append(img_bytes)
            except Exception as e:
                logger.warning(f"Failed to decode image string for PDF report: {e}")

    # Priority 2: Fallback to high-DPI 300 DPI Matplotlib Engine if raw browser images absent
    if not chart_bytes_list and visualizations and len(visualizations) > 0:
        for viz in visualizations:
            try:
                cs = _render_chart_spec_to_png(viz)
                if cs:
                    chart_bytes_list.append(cs.getvalue())
            except Exception as e:
                logger.warning(f"Failed to render chart stream for PDF report: {e}")

    if chart_bytes_list:
        story.append(Spacer(1, 16))
        story.append(Paragraph("Executive Visualizations & Key Charts", h1_style))
        story.append(Spacer(1, 12))

        for img_b in chart_bytes_list:
            try:
                stream = io.BytesIO(img_b)
                img = RLImage(stream, width=5.8 * inch, height=2.9 * inch)
                img.hAlign = 'CENTER'
                story.append(img)
                story.append(Spacer(1, 14))
            except Exception as e:
                logger.warning(f"Failed to add picture to PDF report: {e}")

    doc.build(story, canvasmaker=NumberedCanvas)
    buffer.seek(0)
    return buffer.getvalue()
